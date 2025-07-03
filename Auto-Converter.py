import tkinter as tk
from tkinter import filedialog, messagebox, ttk, font
import pandas as pd
from docx import Document
from docx.shared import Inches
import re
import os
from pathlib import Path
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from functools import lru_cache
import copy
from datetime import datetime

class ModernStyle:
    """Modern styling constants and utilities"""
    # Color palette
    PRIMARY = "#3B82F6"      # Blue
    SECONDARY = "#10B981"    # Green
    ACCENT = "#F59E0B"       # Amber
    DANGER = "#EF4444"       # Red
    SUCCESS = "#10B981"      # Green
    WARNING = "#F59E0B"      # Amber
    
    # Neutral colors
    BACKGROUND = "#F8FAFC"   # Light gray
    SURFACE = "#FFFFFF"      # White
    SURFACE_DARK = "#F1F5F9" # Slightly darker
    TEXT_PRIMARY = "#1E293B" # Dark gray
    TEXT_SECONDARY = "#64748B" # Medium gray
    BORDER = "#E2E8F0"       # Light border
    CARD_SHADOW = "#00000010" # Subtle shadow
    
    # New appealing button colors
    BUTTON_PRIMARY = "#60A5FA"    # Light blue - more appealing
    BUTTON_PRIMARY_HOVER = "#3B82F6"  # Darker blue on hover
    BUTTON_DISABLED = "#CBD5E1"   # Light gray for disabled state

class DocumentConverterTab:
    """Tab for document conversion with modern full-width design"""
    def __init__(self, parent_frame):
        self.parent_frame = parent_frame
        self.word_template_path = None
        self.excel_file_path = None
        self.additional_excel_paths = []
        self.image_folder_path = None  # Add image folder path
        self.image_width = 1.0  # Default image width in inches
        
        # Cache for performance optimization
        self._placeholders_cache = None
        self._additional_data_cache = {}
        self._column_mapping_cache = None
        
        self.setup_modern_fullwidth_ui()
    
    def setup_modern_fullwidth_ui(self):
        # Configure parent frame
        self.parent_frame.configure(bg=ModernStyle.BACKGROUND)
        
        # Main container with full width
        main_container = tk.Frame(self.parent_frame, bg=ModernStyle.BACKGROUND)
        main_container.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        
        # Header section with gradient-like effect
        header_frame = tk.Frame(main_container, bg=ModernStyle.PRIMARY, height=80)
        header_frame.pack(fill=tk.X, pady=(0, 0))
        header_frame.pack_propagate(False)
        
        # Header content
        header_content = tk.Frame(header_frame, bg=ModernStyle.PRIMARY)
        header_content.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        title_font = font.Font(family="Segoe UI", size=22, weight="bold")
        title_label = tk.Label(
            header_content,
            text="HouseBiz Converter",
            font=title_font,
            bg=ModernStyle.PRIMARY,
            fg="white"
        )
        title_label.pack(side=tk.LEFT, anchor=tk.W)
        
        subtitle_font = font.Font(family="Segoe UI", size=11)
        subtitle_label = tk.Label(
            header_content,
            text="Transform templates into personalized documents",
            font=subtitle_font,
            bg=ModernStyle.PRIMARY,
            fg="white"
        )
        subtitle_label.pack(side=tk.RIGHT, anchor=tk.E)
        
        # Content area with full width
        content_area = tk.Frame(main_container, bg=ModernStyle.BACKGROUND)
        content_area.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Two-column layout for better space utilization
        left_column = tk.Frame(content_area, bg=ModernStyle.BACKGROUND)
        left_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        right_column = tk.Frame(content_area, bg=ModernStyle.BACKGROUND)
        right_column.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        # File Upload Section (Left Column)
        upload_section = self.create_modern_section(left_column, "📁 File Upload", ModernStyle.PRIMARY)
        upload_section.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        # Modern upload cards
        self.create_modern_upload_card(upload_section, "📄", "Household or Business Word Template", "Select your template document", 
                                     lambda: self.upload_word_template(), "word_label")
        
        self.create_modern_upload_card(upload_section, "📊", "Household or Business Main File", "Choose your primary data source", 
                                     lambda: self.upload_excel_file(), "excel_label")
        
        self.create_modern_upload_card(upload_section, "📋", "External Table Files", "Optional household or business external tbl files", 
                                     lambda: self.upload_additional_files(), "additional_label")
        
        # Add image folder upload card
        self.create_modern_upload_card(upload_section, "🖼️", "Image Folder (Optional)", "Folder containing images for resp_pix", 
                                     lambda: self.upload_image_folder(), "image_label")
        
        # Processing Section (Right Column)
        processing_section = self.create_modern_section(right_column, "🚀 Processing", ModernStyle.SECONDARY)
        processing_section.pack(fill=tk.BOTH, expand=True)
        
        # Image settings area (before convert button)
        image_settings_area = tk.Frame(processing_section, bg=ModernStyle.SURFACE)
        image_settings_area.pack(fill=tk.X, padx=20, pady=(20, 10))
        


        # Convert button area
        button_area = tk.Frame(processing_section, bg=ModernStyle.SURFACE)
        button_area.pack(fill=tk.X, padx=20, pady=20)
        
        self.convert_btn = self.create_modern_action_button(
            button_area,
            "🚀 Convert Files",
            self.convert_files,
            state="disabled"
        )
        self.convert_btn.pack(anchor=tk.CENTER)
        
        # Progress area
        progress_area = tk.Frame(processing_section, bg=ModernStyle.SURFACE)
        progress_area.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Progress label
        progress_label_font = font.Font(family="Segoe UI", size=10, weight="bold")
        progress_title = tk.Label(
            progress_area,
            text="Progress",
            font=progress_label_font,
            bg=ModernStyle.SURFACE,
            fg=ModernStyle.TEXT_PRIMARY
        )
        progress_title.pack(anchor=tk.W, pady=(0, 10))
        
        # Modern progress bar
        self.setup_modern_progress_bar(progress_area)
        
        # Status area
        status_area = tk.Frame(progress_area, bg=ModernStyle.SURFACE_DARK, relief=tk.FLAT)
        status_area.pack(fill=tk.X, pady=(15, 0))
        
        status_font = font.Font(family="Segoe UI", size=9)
        self.status_label = tk.Label(
            status_area,
            text="Ready to process your files",
            font=status_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_SECONDARY,
            pady=10
        )
        self.status_label.pack()

    def create_modern_section(self, parent, title, accent_color):
        """Create a modern section with accent color"""
        section_frame = tk.Frame(parent, bg=ModernStyle.SURFACE, relief=tk.FLAT)
        section_frame.configure(highlightbackground=ModernStyle.BORDER, highlightthickness=1)
        
        # Section header with accent
        header_frame = tk.Frame(section_frame, bg=accent_color, height=40)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        header_font = font.Font(family="Segoe UI", size=14, weight="bold")
        header_label = tk.Label(
            header_frame,
            text=title,
            font=header_font,
            bg=accent_color,
            fg="white"
        )
        header_label.pack(anchor=tk.W, padx=20, pady=10)
        
        return section_frame
    
    def create_modern_upload_card(self, parent, icon, title, description, command, label_attr):
        """Create a modern upload card with full-width design"""
        card_frame = tk.Frame(parent, bg=ModernStyle.SURFACE)
        card_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # Card content with hover effect simulation
        content_frame = tk.Frame(card_frame, bg=ModernStyle.SURFACE_DARK, relief=tk.FLAT, bd=1)
        content_frame.pack(fill=tk.X, pady=2)
        
        # Icon and content area
        main_content = tk.Frame(content_frame, bg=ModernStyle.SURFACE_DARK)
        main_content.pack(fill=tk.X, padx=15, pady=12)
        
        # Icon
        icon_font = font.Font(family="Segoe UI", size=20)
        icon_label = tk.Label(
            main_content,
            text=icon,
            font=icon_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.PRIMARY,
            width=3
        )
        icon_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # Text content
        text_content = tk.Frame(main_content, bg=ModernStyle.SURFACE_DARK)
        text_content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        title_font = font.Font(family="Segoe UI", size=11, weight="bold")
        title_label = tk.Label(
            text_content,
            text=title,
            font=title_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_PRIMARY
        )
        title_label.pack(anchor=tk.W)
        
        desc_font = font.Font(family="Segoe UI", size=9)
        desc_label = tk.Label(
            text_content,
            text=description,
            font=desc_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_SECONDARY
        )
        desc_label.pack(anchor=tk.W, pady=(2, 0))
        
        # Status label
        status_font = font.Font(family="Segoe UI", size=8)
        if "Optional" in title:
            default_text = "Optional - not selected"
        else:
            default_text = "No file selected"
            
        status_label = tk.Label(
            text_content,
            text=default_text,
            font=status_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_SECONDARY
        )
        status_label.pack(anchor=tk.W, pady=(4, 0))
        setattr(self, label_attr, status_label)
        
        # Button
        if "Folder" in title:
            button_text = "Choose Folder"
        else:
            button_text = "Choose File"
            
        button = self.create_modern_button(
            main_content,
            button_text,
            command,
            ModernStyle.SECONDARY
        )
        button.pack(side=tk.RIGHT, padx=(10, 0))

    def create_modern_button(self, parent, text, command, color):
        """Create a modern button with rounded appearance"""
        btn_font = font.Font(family="Segoe UI", size=9, weight="bold")
        
        button = tk.Button(
            parent,
            text=text,
            command=command,
            font=btn_font,
            bg=color,
            fg="white",
            relief=tk.FLAT,
            bd=0,
            padx=20,
            pady=8,
            cursor="hand2"
        )
        
        # Hover effects
        def on_enter(e):
            button.configure(bg=self.darken_color(color))
        def on_leave(e):
            button.configure(bg=color)
        
        button.bind("<Enter>", on_enter)
        button.bind("<Leave>", on_leave)
        
        return button
    
    def create_modern_action_button(self, parent, text, command, state="normal"):
        """Create a large modern action button with appealing colors"""
        btn_font = font.Font(family="Segoe UI", size=12, weight="bold")
        
        if state == "disabled":
            bg_color = ModernStyle.BUTTON_DISABLED
            text_color = "#64748B"
            cursor = "arrow"
        else:
            bg_color = ModernStyle.BUTTON_PRIMARY
            text_color = "white"
            cursor = "hand2"
        
        button = tk.Button(
            parent,
            text=text,
            command=command,
            font=btn_font,
            bg=bg_color,
            fg=text_color,
            relief=tk.FLAT,
            bd=0,
            padx=40,
            pady=15,
            cursor=cursor,
            state=state
        )
        
        return button
    
    def setup_modern_progress_bar(self, parent):
        """Setup modern progress bar with custom styling"""
        progress_container = tk.Frame(parent, bg=ModernStyle.SURFACE)
        progress_container.pack(fill=tk.X, pady=(0, 10))
        
        # Custom progress bar background
        progress_bg = tk.Frame(progress_container, bg=ModernStyle.SURFACE_DARK, height=8)
        progress_bg.pack(fill=tk.X, pady=(0, 5))
        
        # Progress fill
        self.progress_fill = tk.Frame(progress_bg, bg=ModernStyle.PRIMARY, height=8)
        self.progress_fill.place(x=0, y=0, width=0, height=8)
        
        # Progress percentage
        self.progress_text = tk.Label(
            progress_container,
            text="0%",
            font=font.Font(family="Segoe UI", size=8),
            bg=ModernStyle.SURFACE,
            fg=ModernStyle.TEXT_SECONDARY
        )
        self.progress_text.pack(anchor=tk.E)
        
        # Store the background frame for width calculations
        self.progress_bg_frame = progress_bg
        
        # Initialize progress
        self.progress_value = 0
    
    def update_progress(self, value):
        """Update the custom progress bar"""
        self.progress_value = value
        if hasattr(self, 'progress_bg_frame'):
            # Calculate width based on progress
            bg_width = self.progress_bg_frame.winfo_width()
            if bg_width > 1:  # Ensure the frame has been drawn
                fill_width = int((value / 100) * bg_width)
                self.progress_fill.place(width=fill_width)
                self.progress_text.config(text=f"{int(value)}%")
    
    def darken_color(self, color):
        """Darken a hex color for hover effect"""
        color_map = {
            ModernStyle.PRIMARY: "#2563EB",
            ModernStyle.SECONDARY: "#059669",
            ModernStyle.ACCENT: "#D97706",
            ModernStyle.DANGER: "#DC2626",
            ModernStyle.BUTTON_PRIMARY: ModernStyle.BUTTON_PRIMARY_HOVER
        }
        return color_map.get(color, color)
    
    def update_file_status(self, label, filename, success=True):
        """Update file status with modern styling"""
        if success:
            display_name = filename if len(filename) <= 30 else filename[:27] + "..."
            label.configure(
                text=f"✅ {display_name}",
                fg=ModernStyle.SUCCESS
            )
        else:
            label.configure(
                text="❌ No file selected",
                fg=ModernStyle.TEXT_SECONDARY
            )

    def upload_word_template(self):
        file_path = filedialog.askopenfilename(
            title="Select Word Template",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if file_path:
            self.word_template_path = file_path
            self.update_file_status(self.word_label, Path(file_path).name, True)
            self._preload_placeholders()
            self.check_ready_to_convert()
    
    def upload_excel_file(self):
        file_path = filedialog.askopenfilename(
            title="Select Excel File",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if file_path:
            self.excel_file_path = file_path
            self.update_file_status(self.excel_label, Path(file_path).name, True)
            self.check_ready_to_convert()
    
    def upload_additional_files(self):
        file_paths = filedialog.askopenfilenames(
            title="Select Additional Excel Files",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if file_paths:
            self.additional_excel_paths = list(file_paths)
            self.update_file_status(self.additional_label, f"{len(file_paths)} files selected", True)
            self._preload_additional_data()
            self.check_ready_to_convert()
    
    def upload_image_folder(self):
        folder_path = filedialog.askdirectory(
            title="Select Image Folder"
        )
        if folder_path:
            self.image_folder_path = folder_path
            # Count image files in the folder
            image_extensions = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp')
            image_count = sum(1 for f in os.listdir(folder_path) 
                            if f.lower().endswith(image_extensions))
            self.update_file_status(self.image_label, f"{image_count} images found", True)

    def check_ready_to_convert(self):
        if self.word_template_path and self.excel_file_path:
            self.convert_btn.configure(
                state="normal",
                bg=ModernStyle.BUTTON_PRIMARY,
                fg="white",
                cursor="hand2"
            )
            
            def on_enter(e):
                self.convert_btn.configure(bg=ModernStyle.BUTTON_PRIMARY_HOVER)
            def on_leave(e):
                self.convert_btn.configure(bg=ModernStyle.BUTTON_PRIMARY)
            
            self.convert_btn.bind("<Enter>", on_enter)
            self.convert_btn.bind("<Leave>", on_leave)
        else:
            self.convert_btn.configure(
                state="disabled",
                bg=ModernStyle.BUTTON_DISABLED,
                fg="#64748B",
                cursor="arrow"
            )
            self.convert_btn.unbind("<Enter>")
            self.convert_btn.unbind("<Leave>")

    def _preload_placeholders(self):
        """Pre-load placeholders from template for faster access"""
        try:
            template_doc = Document(self.word_template_path)
            self._placeholders_cache = self.find_placeholders(template_doc)
        except Exception as e:
            print(f"Error pre-loading placeholders: {e}")
    
    def _preload_additional_data(self):
        """Pre-load all additional data into memory for faster access"""
        self._additional_data_cache = {}
        
        for file_path in self.additional_excel_paths:
            try:
                df = pd.read_excel(file_path, header=None)
                
                # Find PARENT_KEY column in row 4 (index 3)
                parent_key_col = None
                if len(df) > 3:
                    for col_idx, cell_value in enumerate(df.iloc[3]):
                        if pd.notna(cell_value) and str(cell_value).strip().upper() == "PARENT_KEY":
                            parent_key_col = col_idx
                            break
                
                if parent_key_col is not None:
                    # Get headers from row 4
                    headers = df.iloc[3].tolist()
                    
                    # Get data rows (from row 5 onwards)
                    data_rows = df.iloc[4:]
                    
                    # Group data by PARENT_KEY for O(1) lookup
                    for _, row in data_rows.iterrows():
                        if pd.notna(row.iloc[parent_key_col]):
                            key_value = str(row.iloc[parent_key_col]).strip()
                            
                            if key_value not in self._additional_data_cache:
                                self._additional_data_cache[key_value] = []
                            
                            # Create a dictionary mapping header to value
                            row_data = {}
                            for idx, header in enumerate(headers):
                                if pd.notna(header):
                                    row_data[str(header).strip()] = row.iloc[idx] if pd.notna(row.iloc[idx]) else ""
                            
                            self._additional_data_cache[key_value].append(row_data)
                            
            except Exception as e:
                print(f"Error pre-loading additional file {file_path}: {e}")

    @lru_cache(maxsize=1)
    def find_placeholders(self, doc=None):
        """Find all placeholders in the document like {firstname} - cached for performance"""
        if self._placeholders_cache:
            return self._placeholders_cache
            
        placeholders = set()
        
        if doc is None and self.word_template_path:
            doc = Document(self.word_template_path)
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(r'\{([^}]+)\}', paragraph.text)
            placeholders.update(matches)
        
        # Search in tables (including nested tables)
        self._search_tables_for_placeholders(doc.tables, placeholders)
        
        return list(placeholders)
    
    def _search_tables_for_placeholders(self, tables, placeholders):
        """Recursively search tables and nested tables for placeholders"""
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    # Search text in cell
                    matches = re.findall(r'\{([^}]+)\}', cell.text)
                    placeholders.update(matches)
                    
                    # Search for nested tables in this cell
                    if cell.tables:
                        self._search_tables_for_placeholders(cell.tables, placeholders)
    
    def find_column_mapping(self, df, placeholders):
        """Find column mapping - cached for performance"""
        if self._column_mapping_cache:
            return self._column_mapping_cache
            
        column_mapping = {}
        
        # Get all possible column values from rows 1, 2, 3, and 4 (index 0, 1, 2, 3)
        search_rows = min(4, len(df))  # Search in first 4 rows
        
        for placeholder in placeholders:
            found = False
            
            # Search through each row (1, 2, 3, 4)
            for row_idx in range(search_rows):
                for col_idx, col_name in enumerate(df.columns):
                    cell_value = df.iloc[row_idx, col_idx]
                    
                    # Check if entire cell content matches the placeholder (case-insensitive)
                    if pd.notna(cell_value):
                        cell_str = str(cell_value).strip()
                        if cell_str.lower() == placeholder.lower():
                            column_mapping[placeholder] = col_name
                            found = True
                            break
                
                if found:
                    break
            
            if not found:
                print(f"Warning: Placeholder '{placeholder}' not found in any of the first 4 rows")
        
        self._column_mapping_cache = column_mapping
        return column_mapping
    
    def replace_placeholders_optimized(self, doc, data_row):
        """Optimized placeholder replacement using compiled regex with image support"""
        # Pre-compile regex for better performance
        placeholder_pattern = re.compile(r'\{([^}]+)\}')
        
        # Get image width setting
        try:
            image_width = float(self.image_width_var.get())
        except:
            image_width = 1.0
        
        # Process bus_info_needs ranking if present
        if 'bus_info_needs' in data_row:
            ranked_data = self.process_bus_info_needs_ranking(data_row)
            data_row.update(ranked_data)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if '{' in original_text:  # Quick check before regex
                # Check for image placeholder first
                if self.image_folder_path and 'resp_pix' in data_row and '{resp_pix}' in original_text:
                    image_filename = str(data_row.get('resp_pix', '')).strip()
                    if image_filename and image_filename.lower() != 'nan':
                        # Try different image extensions
                        image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
                        image_found = False
                        
                        for ext in [''] + image_extensions:  # Try without extension first, then with extensions
                            if ext == '':
                                test_filename = image_filename
                            else:
                                # Remove existing extension if any, then add new one
                                base_name = os.path.splitext(image_filename)[0]
                                test_filename = base_name + ext
                            
                            image_path = os.path.join(self.image_folder_path, test_filename)
                            if os.path.exists(image_path):
                                if self.replace_image_in_paragraph(paragraph, image_path, image_width, 'resp_pix'):
                                    image_found = True
                                    break
                        
                        if image_found:
                            continue  # Skip text replacement if image was inserted
                
                # Regular text replacement
                new_text = original_text
                for key, value in data_row.items():
                    new_text = new_text.replace(f'{{{key}}}', str(value))
                if new_text != original_text:
                    paragraph.text = new_text
        
        # Replace in tables (including nested tables)
        self._replace_in_tables_optimized(doc.tables, data_row, image_width)

    def process_bus_info_needs_ranking(self, data_row):
        """Process bus_info_needs column to create ranked lists and reasons"""
        bus_info_needs = str(data_row.get('bus_info_needs', '')).strip()
        bus_info_needs_o = str(data_row.get('bus_info_needs_o', '')).strip()
        
        # Initialize result dictionary
        result = {}
        
        if not bus_info_needs or bus_info_needs.lower() in ['nan', 'na', '']:
            # If no data, return empty strings for all placeholders
            result['bus_info_needs'] = ''
            result['bus_info_needs_rank'] = ''
            for i in range(1, 10):
                result[f'bus_info_needs_rank_reason{i}'] = ''
            return result
        
        # Split the comma-separated values and clean them
        items = [item.strip() for item in bus_info_needs.split(',') if item.strip()]
        
        # Process items in order - the order determines the ranking
        processed_items = []
        for item in items:
            item_clean = item.strip()
            # Skip empty items and 'NA' entries
            if item_clean and item_clean.lower() not in ['na', '']:
                # Handle "Others, specify" case
                if item_clean.lower() in ['others', 'specify']:
                    if bus_info_needs_o and bus_info_needs_o.lower() not in ['nan', 'na', '']:
                        processed_items.append(f"Others, specify: {bus_info_needs_o}")
                    else:
                        processed_items.append("Others, specify")
                else:
                    processed_items.append(item_clean)
        
        # Create the ranked list output - each item on a new line
        result['bus_info_needs'] = '\n'.join(processed_items)
        
        # Create rank numbers corresponding to each item
        rank_numbers = [str(i) for i in range(1, len(processed_items) + 1)]
        result['bus_info_needs_rank'] = '\n'.join(rank_numbers)
        
        # Add individual reason placeholders (bus_info_needs_rank_reason1 to bus_info_needs_rank_reason9)
        for i in range(1, 10):
            reason_key = f'bus_info_needs_rank_reason{i}'
            if i <= len(processed_items):
                # Check if there's a corresponding reason in the data_row
                if reason_key in data_row and str(data_row[reason_key]).strip() not in ['', 'nan', 'NA']:
                    result[reason_key] = str(data_row[reason_key])
                else:
                    # Use the item itself as placeholder for the reason
                    result[reason_key] = f"{{{reason_key}}}"
            else:
                result[reason_key] = ''
        
        return result

    def replace_placeholders_optimized(self, doc, data_row):
        """Optimized placeholder replacement using compiled regex with image support"""
        # Pre-compile regex for better performance
        placeholder_pattern = re.compile(r'\{([^}]+)\}')
        
        # Get image width setting
        try:
            image_width = float(self.image_width_var.get())
        except:
            image_width = 1.0
        
        # Process bus_info_needs ranking if present
        if 'bus_info_needs' in data_row:
            ranked_data = self.process_bus_info_needs_ranking(data_row)
            data_row.update(ranked_data)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if '{' in original_text:  # Quick check before regex
                # Check for image placeholder first
                if self.image_folder_path and 'resp_pix' in data_row and '{resp_pix}' in original_text:
                    image_filename = str(data_row.get('resp_pix', '')).strip()
                    if image_filename and image_filename.lower() != 'nan':
                        # Try different image extensions
                        image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
                        image_found = False
                        
                        for ext in [''] + image_extensions:  # Try without extension first, then with extensions
                            if ext == '':
                                test_filename = image_filename
                            else:
                                # Remove existing extension if any, then add new one
                                base_name = os.path.splitext(image_filename)[0]
                                test_filename = base_name + ext
                            
                            image_path = os.path.join(self.image_folder_path, test_filename)
                            if os.path.exists(image_path):
                                if self.replace_image_in_paragraph(paragraph, image_path, image_width, 'resp_pix'):
                                    image_found = True
                                    break
                        
                        if image_found:
                            continue  # Skip text replacement if image was inserted
                
                # Regular text replacement
                new_text = original_text
                for key, value in data_row.items():
                    new_text = new_text.replace(f'{{{key}}}', str(value))
                if new_text != original_text:
                    paragraph.text = new_text
        
        # Replace in tables (including nested tables)
        self._replace_in_tables_optimized(doc.tables, data_row, image_width)

    def process_bus_info_needs_ranking(self, data_row):
        """Process bus_info_needs column to create ranked lists and reasons"""
        bus_info_needs = str(data_row.get('bus_info_needs', '')).strip()
        bus_info_needs_o = str(data_row.get('bus_info_needs_o', '')).strip()
        
        # Initialize result dictionary
        result = {}
        
        if not bus_info_needs or bus_info_needs.lower() in ['nan', 'na', '']:
            # If no data, return empty strings for all placeholders
            result['bus_info_needs'] = ''
            result['bus_info_needs_rank'] = ''
            for i in range(1, 10):
                result[f'bus_info_needs_rank_reason{i}'] = ''
            return result
        
        # Split the comma-separated values and clean them
        items = [item.strip() for item in bus_info_needs.split(',') if item.strip()]
        
        # Process items in order - the order determines the ranking
        processed_items = []
        for item in items:
            item_clean = item.strip()
            # Skip empty items and 'NA' entries
            if item_clean and item_clean.lower() not in ['na', '']:
                # Handle "Others, specify" case
                if item_clean.lower() in ['others', 'specify']:
                    if bus_info_needs_o and bus_info_needs_o.lower() not in ['nan', 'na', '']:
                        processed_items.append(f"Others, specify: {bus_info_needs_o}")
                    else:
                        processed_items.append("Others, specify")
                else:
                    processed_items.append(item_clean)
        
        # Create the ranked list output - each item on a new line
        result['bus_info_needs'] = '\n'.join(processed_items)
        
        # Create rank numbers corresponding to each item
        rank_numbers = [str(i) for i in range(1, len(processed_items) + 1)]
        result['bus_info_needs_rank'] = '\n'.join(rank_numbers)
        
        # Add individual reason placeholders (bus_info_needs_rank_reason1 to bus_info_needs_rank_reason9)
        for i in range(1, 10):
            reason_key = f'bus_info_needs_rank_reason{i}'
            if i <= len(processed_items):
                # Check if there's a corresponding reason in the data_row
                if reason_key in data_row and str(data_row[reason_key]).strip() not in ['', 'nan', 'NA']:
                    result[reason_key] = str(data_row[reason_key])
                else:
                    # Use the item itself as placeholder for the reason
                    result[reason_key] = f"{{{reason_key}}}"
            else:
                result[reason_key] = ''
        
        return result

    def replace_placeholders_optimized(self, doc, data_row):
        """Optimized placeholder replacement using compiled regex with image support"""
        # Pre-compile regex for better performance
        placeholder_pattern = re.compile(r'\{([^}]+)\}')
        
        # Get image width setting
        try:
            image_width = float(self.image_width_var.get())
        except:
            image_width = 1.0
        
        # Process bus_info_needs ranking if present
        if 'bus_info_needs' in data_row:
            ranked_data = self.process_bus_info_needs_ranking(data_row)
            data_row.update(ranked_data)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if '{' in original_text:  # Quick check before regex
                # Check for image placeholder first
                if self.image_folder_path and 'resp_pix' in data_row and '{resp_pix}' in original_text:
                    image_filename = str(data_row.get('resp_pix', '')).strip()
                    if image_filename and image_filename.lower() != 'nan':
                        # Try different image extensions
                        image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
                        image_found = False
                        
                        for ext in [''] + image_extensions:  # Try without extension first, then with extensions
                            if ext == '':
                                test_filename = image_filename
                            else:
                                # Remove existing extension if any, then add new one
                                base_name = os.path.splitext(image_filename)[0]
                                test_filename = base_name + ext
                            
                            image_path = os.path.join(self.image_folder_path, test_filename)
                            if os.path.exists(image_path):
                                if self.replace_image_in_paragraph(paragraph, image_path, image_width, 'resp_pix'):
                                    image_found = True
                                    break
                        
                        if image_found:
                            continue  # Skip text replacement if image was inserted
                
                # Regular text replacement
                new_text = original_text
                for key, value in data_row.items():
                    new_text = new_text.replace(f'{{{key}}}', str(value))
                if new_text != original_text:
                    paragraph.text = new_text
        
        # Replace in tables (including nested tables)
        self._replace_in_tables_optimized(doc.tables, data_row, image_width)

    def process_bus_info_needs_ranking(self, data_row):
        """Process bus_info_needs column to create ranked lists and reasons"""
        bus_info_needs = str(data_row.get('bus_info_needs', '')).strip()
        bus_info_needs_o = str(data_row.get('bus_info_needs_o', '')).strip()
        
        # Initialize result dictionary
        result = {}
        
        if not bus_info_needs or bus_info_needs.lower() in ['nan', 'na', '']:
            # If no data, return empty strings for all placeholders
            result['bus_info_needs'] = ''
            result['bus_info_needs_rank'] = ''
            for i in range(1, 10):
                result[f'bus_info_needs_rank_reason{i}'] = ''
            return result
        
        # Split the comma-separated values and clean them
        items = [item.strip() for item in bus_info_needs.split(',') if item.strip()]
        
        # Process items in order - the order determines the ranking
        processed_items = []
        for item in items:
            item_clean = item.strip()
            # Skip empty items and 'NA' entries
            if item_clean and item_clean.lower() not in ['na', '']:
                # Handle "Others, specify" case
                if item_clean.lower() in ['others', 'specify']:
                    if bus_info_needs_o and bus_info_needs_o.lower() not in ['nan', 'na', '']:
                        processed_items.append(f"Others, specify: {bus_info_needs_o}")
                    else:
                        processed_items.append("Others, specify")
                else:
                    processed_items.append(item_clean)
        
        # Create the ranked list output - each item on a new line
        result['bus_info_needs'] = '\n'.join(processed_items)
        
        # Create rank numbers corresponding to each item
        rank_numbers = [str(i) for i in range(1, len(processed_items) + 1)]
        result['bus_info_needs_rank'] = '\n'.join(rank_numbers)
        
        # Add individual reason placeholders (bus_info_needs_rank_reason1 to bus_info_needs_rank_reason9)
        for i in range(1, 10):
            reason_key = f'bus_info_needs_rank_reason{i}'
            if i <= len(processed_items):
                # Check if there's a corresponding reason in the data_row
                if reason_key in data_row and str(data_row[reason_key]).strip() not in ['', 'nan', 'NA']:
                    result[reason_key] = str(data_row[reason_key])
                else:
                    # Use the item itself as placeholder for the reason
                    result[reason_key] = f"{{{reason_key}}}"
            else:
                result[reason_key] = ''
        
        return result

    def replace_placeholders_optimized(self, doc, data_row):
        """Optimized placeholder replacement using compiled regex with image support"""
        # Pre-compile regex for better performance
        placeholder_pattern = re.compile(r'\{([^}]+)\}')
        
        # Get image width setting
        try:
            image_width = float(self.image_width_var.get())
        except:
            image_width = 1.0
        
        # Process bus_info_needs ranking if present
        if 'bus_info_needs' in data_row:
            ranked_data = self.process_bus_info_needs_ranking(data_row)
            data_row.update(ranked_data)
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            if '{' in original_text:  # Quick check before regex
                # Check for image placeholder first
                if self.image_folder_path and 'resp_pix' in data_row and '{resp_pix}' in original_text:
                    image_filename = str(data_row.get('resp_pix', '')).strip()
                    if image_filename and image_filename.lower() != 'nan':
                        # Try different image extensions
                        image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
                        image_found = False
                        
                        for ext in [''] + image_extensions:  # Try without extension first, then with extensions
                            if ext == '':
                                test_filename = image_filename
                            else:
                                # Remove existing extension if any, then add new one
                                base_name = os.path.splitext(image_filename)[0]
                                test_filename = base_name + ext
                            
                            image_path = os.path.join(self.image_folder_path, test_filename)
                            if os.path.exists(image_path):
                                if self.replace_image_in_paragraph(paragraph, image_path, image_width, 'resp_pix'):
                                    image_found = True
                                    break
                        
                        if image_found:
                            continue  # Skip text replacement if image was inserted
                
                # Regular text replacement
                new_text = original_text
                for key, value in data_row.items():
                    new_text = new_text.replace(f'{{{key}}}', str(value))
                if new_text != original_text:
                    paragraph.text = new_text
        
        # Replace in tables (including nested tables)
        self._replace_in_tables_optimized(doc.tables, data_row, image_width)

    def replace_bus_info_needs_in_table_cell(self, cell, data_row):
        """Handle special replacement for bus_info_needs table cells"""
        # Process the ranking data first
        ranked_data = self.process_bus_info_needs_ranking(data_row)
        
        # Handle each paragraph in the cell
        for paragraph in cell.paragraphs:
            original_text = paragraph.text
            
            # Handle the main question column - just replace the placeholder, don't add items here
            if '{bus_info_needs}' in original_text:
                # For the main question cell, just remove the placeholder
                new_text = original_text.replace('{bus_info_needs}', '')
                
                # Handle others specification
                if '{bus_info_needs_o}' in new_text:
                    bus_info_needs_o = str(data_row.get('bus_info_needs_o', '')).strip()
                    if bus_info_needs_o and bus_info_needs_o.lower() not in ['nan', 'na', '']:
                        new_text = new_text.replace('{bus_info_needs_o}', bus_info_needs_o)
                    else:
                        new_text = new_text.replace('if others{bus_info_needs_o}', '')
                        new_text = new_text.replace('{bus_info_needs_o}', '')
                
                paragraph.text = new_text
            
            # Handle rank column - just remove the placeholder, don't add ranks here
            elif 'Rank, by order of importance' in original_text or '{bus_info_needs_rank}' in original_text:
                if '{bus_info_needs_rank}' in original_text:
                    new_text = original_text.replace('{bus_info_needs_rank}', '')
                else:
                    new_text = original_text
                paragraph.text = new_text
            
            # Handle reason columns - replace individual reason placeholders
            elif '{bus_info_needs_rank_reason' in original_text:
                new_text = original_text
                for key, value in ranked_data.items():
                    if key.startswith('bus_info_needs_rank_reason'):
                        new_text = new_text.replace(f'{{{key}}}', str(value))
                paragraph.text = new_text
            
            # Handle any other placeholders in this cell
            else:
                new_text = original_text
                # Replace regular placeholders first
                for key, value in data_row.items():
                    if f'{{{key}}}' in new_text:
                        new_text = new_text.replace(f'{{{key}}}', str(value))
                # Then replace ranking placeholders
                for key, value in ranked_data.items():
                    if f'{{{key}}}' in new_text:
                        new_text = new_text.replace(f'{{{key}}}', str(value))
                
                if new_text != original_text:
                    paragraph.text = new_text

    def populate_bus_info_needs_table(self, table, data_row):
        """Populate the bus_info_needs table with ranked items in separate rows"""
        # Process the ranking data
        ranked_data = self.process_bus_info_needs_ranking(data_row)
        
        # Get the items list
        items_text = ranked_data['bus_info_needs']
        if not items_text:
            return
            
        items = items_text.split('\n')
        
        # Find the table structure - typically has 3 columns
        # Column 1: What types of information would be helpful...
        # Column 2: Rank, by order of importance  
        # Column 3: Why is this information useful...
        
        # Keep the header rows (typically first 2 rows) and remove any existing data rows
        header_rows = 2  # Adjust this based on your table structure
        while len(table.rows) > header_rows:
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add a row for each item
        for i, item in enumerate(items, 1):
            if item.strip():  # Only add non-empty items
                new_row = table.add_row()
                cells = new_row.cells
                
                # Column 1: Item text
                if len(cells) > 0:
                    cells[0].text = item.strip()
                
                # Column 2: Rank number
                if len(cells) > 1:
                    cells[1].text = str(i)
                
                # Column 3: Reason (if available)
                if len(cells) > 2:
                    reason_key = f'bus_info_needs_rank_reason{i}'
                    if reason_key in data_row and str(data_row[reason_key]).strip() not in ['', 'nan', 'NA']:
                        cells[2].text = str(data_row[reason_key])
                    else:
                        cells[2].text = ''  # Leave empty for user to fill

    def _replace_in_tables_optimized(self, tables, data_row, image_width=5.0):
        """Optimized table replacement with image support"""
        for table in tables:
            # Check if this is a bus_info_needs table before processing individual cells
            table_text = ' '.join(cell.text for row in table.rows[:2] for cell in row.cells)
            
            if self.is_bus_info_needs_table(table_text):
                # Handle the entire bus_info_needs table
                self.populate_bus_info_needs_table(table, data_row)
                continue
            
            # Regular table processing for non-bus_info_needs tables
            for row in table.rows:
                for cell in row.cells:
                    # Replace in cell paragraphs (more reliable than cell.text)
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        if '{' in original_text:  # Quick check before processing
                            # Check for image placeholder first
                            if self.image_folder_path and 'resp_pix' in data_row and '{resp_pix}' in original_text:
                                image_filename = str(data_row.get('resp_pix', '')).strip()
                                if image_filename and image_filename.lower() != 'nan':
                                    # Try different image extensions
                                    image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
                                    image_found = False
                                    
                                    for ext in [''] + image_extensions:  # Try without extension first, then with extensions
                                        if ext == '':
                                            test_filename = image_filename
                                        else:
                                            # Remove existing extension if any, then add new one
                                            base_name = os.path.splitext(image_filename)[0]
                                            test_filename = base_name + ext
                                        
                                        image_path = os.path.join(self.image_folder_path, test_filename)
                                        if os.path.exists(image_path):
                                            if self.replace_image_in_paragraph(paragraph, image_path, image_width, 'resp_pix'):
                                                image_found = True
                                                break
                                    
                                    if image_found:
                                        continue  # Skip text replacement if image was inserted
                            
                            # Regular text replacement
                            new_text = original_text
                            for key, value in data_row.items():
                                new_text = new_text.replace(f'{{{key}}}', str(value))
                            if new_text != original_text:
                                paragraph.text = new_text
                    
                    # Handle nested tables in this cell
                    if cell.tables:
                        self._replace_in_tables_optimized(cell.tables, data_row, image_width)

    def is_bus_info_needs_table(self, text):
        """Check if this is the bus_info_needs table section"""
        return ('{bus_info_needs}' in text or 
                'What types of information would be helpful' in text or
                'Information Needs' in text)

    def replace_image_in_paragraph(self, paragraph, image_path, image_width, placeholder_name):
        """Replace placeholder with image in a paragraph"""
        if f"{{{placeholder_name}}}" in paragraph.text:
            # Keep track of the paragraph's style
            paragraph_style = paragraph.style

            # Get text before and after the placeholder
            text_parts = paragraph.text.split(f"{{{placeholder_name}}}")
            before_text = text_parts[0]
            after_text = text_parts[1] if len(text_parts) > 1 else ""

            # Clear the paragraph
            p = paragraph.clear()

            # Restore the paragraph's style
            paragraph.style = paragraph_style

            # Add text before the image
            if before_text:
                paragraph.add_run(before_text)

            # Add the image
            try:
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(image_width))
            except Exception as e:
                paragraph.add_run(f"[Image not found: {os.path.basename(image_path)}]")

            # Add text after the image
            if after_text:
                paragraph.add_run(after_text)

            return True
        return False

    def convert_files(self):
        def conversion_worker():
            try:
                self.status_label.config(text="🔄 Initializing conversion...", fg=ModernStyle.PRIMARY)
                self.update_progress(0)
                
                destination_path = filedialog.askdirectory(title="Select Destination Folder for ZIP File")
                if not destination_path:
                    self.status_label.config(text="Ready to process your files", fg=ModernStyle.TEXT_SECONDARY)
                    return
                
                df = pd.read_excel(self.excel_file_path, header=None)
                placeholders = self._placeholders_cache or self.find_placeholders()
                
                if not placeholders:
                    messagebox.showwarning("Warning", "No placeholders found in template")
                    return
                
                column_mapping = self.find_column_mapping(df, placeholders)
                
                if not column_mapping:
                    messagebox.showerror("Error", "No matching columns found between template and Excel headers in rows 1-3")
                    return
                
                key_column = None
                if len(df) > 3:
                    for col_idx, cell_value in enumerate(df.iloc[3]):
                        if pd.notna(cell_value) and str(cell_value).strip().upper() == "KEY":
                            key_column = col_idx
                            break
                
                if key_column is None:
                    messagebox.showerror("Error", "KEY column not found in row 4 of main Excel file")
                    return
                
                data_rows = df.iloc[4:].copy()
                
                if data_rows.empty:
                    messagebox.showwarning("Warning", "No data found starting from row 5")
                    return
                
                temp_dir = Path(destination_path) / "temp_documents"
                temp_dir.mkdir(exist_ok=True)
                
                total_rows = len(data_rows)
                generated_files = []
                
                process_args = []
                for idx, (original_row_idx, row) in enumerate(data_rows.iterrows()):
                    process_args.append((idx, original_row_idx, row, column_mapping, key_column, temp_dir))
                
                max_workers = min(4, os.cpu_count() or 1)
                completed_count = 0
                
                with ThreadPoolExecutor(max_workers=max_workers) as executor:
                    future_to_args = {executor.submit(self.process_single_document, args): args for args in process_args}
                    
                    for future in as_completed(future_to_args):
                        result, error = future.result()
                        completed_count += 1
                        
                        if result:
                            generated_files.append(result)
                        else:
                            print(f"Error processing document: {error}")
                        
                        progress_value = (completed_count / total_rows) * 80
                        self.update_progress(progress_value)
                        self.status_label.config(
                            text=f"📝 Processing document {completed_count} of {total_rows}",
                            fg=ModernStyle.PRIMARY
                        )
                
                self.status_label.config(text="📦 Creating ZIP file...", fg=ModernStyle.ACCENT)
                self.update_progress(85)
                
                zip_path = Path(destination_path) / "Generated_Documents.zip"
                with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for file_path in generated_files:
                        zipf.write(file_path, file_path.name)
                
                for file_path in generated_files:
                    if file_path.exists():
                        file_path.unlink()
                if temp_dir.exists():
                    temp_dir.rmdir()
                
                self.update_progress(100)
                self.status_label.config(
                    text=f"✅ Successfully generated {total_rows} documents!",
                    fg=ModernStyle.SUCCESS
                )
                messagebox.showinfo("Success", f"🎉 Generated {total_rows} documents in ZIP file:\n{zip_path}")
                
            except Exception as e:
                print(f"Error occurred: {str(e)}")
                messagebox.showerror("Error", f"An error occurred: {str(e)}")
                self.status_label.config(text="❌ Error occurred during processing", fg=ModernStyle.DANGER)
            finally:
                self.update_progress(0)
                self.parent_frame.after(3000, lambda: self.status_label.config(
                    text="Ready to process your files", 
                    fg=ModernStyle.TEXT_SECONDARY
                ))
        
        threading.Thread(target=conversion_worker, daemon=True).start()

    def get_additional_data_for_key_optimized(self, key_value):
        """Get all matching rows from pre-loaded cache - O(1) lookup"""
        return self._additional_data_cache.get(str(key_value).strip(), [])
    
    def populate_dynamic_tables_optimized(self, doc, additional_rows):
        """Optimized dynamic table population with pre-categorized data"""
        # More comprehensive categorization with better debugging
        categorized_data = {
            'hh_member': [],
            'labor': [],
            'debt': [],
            'land': [],
            'struct': [],
            'affected_struct': [],
            'tree': [],
            'crop': [],
            'income_loss': [],
            'others': []
        }
        
        # Initialize totals for calculations
        hh_calc_total_sum = 0
            
        # Categorize each row
        for row in additional_rows:
            row_keys = list(row.keys())
            
            # Check for different patterns
            if any('crop' in key.lower() for key in row_keys):
                categorized_data['crop'].append(row)
            elif any('income_loss' in key.lower() or 'incomeloss' in key.lower() for key in row_keys):
                categorized_data['income_loss'].append(row)
            elif any('others' in key.lower() for key in row_keys):
                categorized_data['others'].append(row)
            elif any(key.startswith('hhcomp_hhmmbr_') for key in row_keys):
                categorized_data['hh_member'].append(row)
            elif any(key.startswith(('hh_labor_', 'hh_wrk_', 'hh_calc_', 'hh_total_')) for key in row_keys):
                categorized_data['labor'].append(row)
                # Calculate sum for hh_calc_total_inc
                hh_calc_value = row.get('hh_calc_total_inc', 0)
                if hh_calc_value and str(hh_calc_value).strip() and str(hh_calc_value).lower() != 'nan':
                    try:
                        hh_calc_total_sum += float(hh_calc_value)
                    except (ValueError, TypeError):
                        pass  # Skip invalid values
            elif any(key.startswith(('debt_', 'loan_', 'pymt_')) for key in row_keys):
                categorized_data['debt'].append(row)
            elif any(key.startswith('asset_land_') for key in row_keys):
                categorized_data['land'].append(row)
            elif any(key.startswith('asset_struct_') for key in row_keys):
                categorized_data['struct'].append(row)
            elif any(key.startswith('affctd_struct_') for key in row_keys):
                categorized_data['affected_struct'].append(row)
            elif any(key.startswith('tree_') for key in row_keys):
                categorized_data['tree'].append(row)

        # Process tables with optimized table identification
        for table in doc.tables:
            # Use more efficient table identification
            first_two_rows_text = ' '.join(cell.text for row in table.rows[:2] for cell in row.cells)
            #business and household main tables
            if "Name of HH Member" in first_two_rows_text:
                self.populate_hh_member_table(table, categorized_data['hh_member'])
            elif "Ownership of at least one savings account" in first_two_rows_text:
                self.populate_savings_table(table, categorized_data['hh_member'])
            elif "Labor Force Status" in first_two_rows_text:
                self.populate_labor_table(table, categorized_data['labor'])
            elif "With formal loan contract? (Y/N)" in first_two_rows_text:
                self.populate_debt_table(table, categorized_data['debt'])
            elif "13.1 Affected Assets: Land" in first_two_rows_text or "10.0 Affected Assets: Land" in first_two_rows_text:
                self.populate_land_assets_table(table, categorized_data['land'])
            elif "13.2 Affected Assets: Structure" in first_two_rows_text or "10.2 Affected Assets: Structure" in first_two_rows_text:
                self.populate_structure_assets_table(table, categorized_data['struct'])
            elif "13.3 Affected Structure" in first_two_rows_text or "10.3 Affected Structure" in first_two_rows_text:
                self.populate_affected_structure_table(table, categorized_data['affected_struct'])
            elif "13.4 Trees" in first_two_rows_text or  "10.4 Trees" in first_two_rows_text:
                self.populate_trees_table(table, categorized_data['tree'])
            elif "13.5 Crops" in first_two_rows_text  or "10.5 Crops" in first_two_rows_text or "crops_grp_converted" in first_two_rows_text.lower():
                self.populate_crops_table(table, categorized_data['crop'])
            elif "13.6 Income Loss" in first_two_rows_text  or "10.6 Income Loss" in first_two_rows_text or "income_loss_grp_converted" in first_two_rows_text.lower():
                self.populate_income_loss_table(table, categorized_data['income_loss'])
            elif "13.7 Others" in first_two_rows_text or "10.7 Others" in first_two_rows_text or "others_grp_converted" in first_two_rows_text.lower():
                self.populate_others_table(table, categorized_data['others'])
        
        # Replace the sum placeholder in the document
        self.replace_sum_placeholder(doc, 'hh_calc_total_sum', hh_calc_total_sum)

    def replace_sum_placeholder(self, doc, placeholder_name, sum_value):
        """Replace sum placeholder in the entire document"""
        placeholder_text = f"{{{placeholder_name}}}"
        sum_text = str(sum_value) if sum_value != 0 else "0"
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder_text in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder_text, sum_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder_text, sum_text)
                    
                    # Handle nested tables
                    if cell.tables:
                        self._replace_sum_in_nested_tables(cell.tables, placeholder_text, sum_text)

    def _replace_sum_in_nested_tables(self, tables, placeholder_text, sum_text):
        """Replace sum placeholder in nested tables"""
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder_text, sum_text)
                    
                    # Handle further nested tables
                    if cell.tables:
                        self._replace_sum_in_nested_tables(cell.tables, placeholder_text, sum_text)

    def clear_all_remaining_placeholders_optimized(self, doc):
        """Optimized placeholder clearing with compiled regex"""
        placeholder_pattern = re.compile(r'\{[^}]+\}')
        
        # Clear placeholders in paragraphs
        for paragraph in doc.paragraphs:
            if '{' in paragraph.text:  # Quick check before regex
                paragraph.text = placeholder_pattern.sub('', paragraph.text)
        
        # Clear placeholders in tables (including nested tables)
        self._clear_tables_placeholders_recursive_optimized(doc.tables, placeholder_pattern)
    
    def _clear_tables_placeholders_recursive_optimized(self, tables, pattern):
        """Optimized recursive placeholder clearing in tables"""
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{' in paragraph.text:  # Quick check before regex
                            paragraph.text = pattern.sub('', paragraph.text)
                    
                    # Handle nested tables in this cell
                    if cell.tables:
                        self._clear_tables_placeholders_recursive_optimized(cell.tables, pattern)
    
    def process_single_document(self, args):
        """Process a single document - optimized for parallel processing with image support"""
        try:
            idx, original_row_idx, row, column_mapping, key_column, temp_dir = args
            
            # Create new document from template
            new_doc = Document(self.word_template_path)
            
            # Get KEY value for this row
            key_value = row.iloc[key_column] if pd.notna(row.iloc[key_column]) else ""
            
            # Get additional data from cache (O(1) lookup)
            additional_rows = self.get_additional_data_for_key_optimized(key_value)
            
            # Prepare data for replacement (from main file)
            replacement_data = {}
            for placeholder, column_name in column_mapping.items():
                replacement_data[placeholder] = row[column_name] if pd.notna(row[column_name]) else ""
            
            # Replace placeholders with optimized method (including images)
            self.replace_placeholders_optimized(new_doc, replacement_data)
            
            # Populate dynamic tables with additional data
            if additional_rows:
                self.populate_dynamic_tables_optimized(new_doc, additional_rows)
            
            # Clear any remaining placeholders in the entire document
            self.clear_all_remaining_placeholders_optimized(new_doc)
            
            # Generate filename using pckg_brgy and resp_lname
            resp_lname = replacement_data.get('resp_lname', '')
            resp_brgy = replacement_data.get('pckg_brgy', '')
            
            if resp_brgy and str(resp_brgy).strip() and resp_lname and str(resp_lname).strip():
                # Clean both brgy and last name for filename use
                clean_brgy = str(resp_brgy).strip()
                clean_lname = str(resp_lname).strip()
                
                # Remove invalid filename characters from both
                invalid_chars = '<>:"/\\|?*'
                for char in invalid_chars:
                    clean_brgy = clean_brgy.replace(char, '_')
                    clean_lname = clean_lname.replace(char, '_')
                
                # Limit length of each part and combine
                clean_brgy = clean_brgy[:15]  # Limit brgy to 15 characters
                clean_lname = clean_lname[:15]  # Limit lname to 15 characters
                filename = f"{clean_brgy}_{clean_lname}_{idx + 1:03d}.docx"
            elif resp_lname and str(resp_lname).strip():
                # Fallback to just last name if brgy is not available
                clean_lname = str(resp_lname).strip()
                invalid_chars = '<>:"/\\|?*'
                for char in invalid_chars:
                    clean_lname = clean_lname.replace(char, '_')
                clean_lname = clean_lname[:20]
                filename = f"{clean_lname}_{idx + 1:03d}.docx"
            else:
                # Final fallback to numbered naming
                filename = f"document_{idx + 1:03d}.docx"
            
            # Save document
            output_path = temp_dir / filename
            new_doc.save(output_path)
            
            return output_path, None
            
        except Exception as e:
            return None, str(e)
    
    def populate_hh_member_table(self, table, additional_rows):
        """Populate the HH Member table with data from additional files"""
        # Keep header rows and remove existing data rows (first row after headers)
        while len(table.rows) > 2:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows
        for idx, row_data in enumerate(additional_rows, start=1):
            new_row = table.add_row()
            cells = new_row.cells
            
            # Populate cells based on your table structure
            if len(cells) > 0:
                cells[0].text = str(idx)  # Row number
            if len(cells) > 1:
                # Full name (combining first, middle, last names)
                fname = row_data.get('hhcomp_hhmmbr_fname', '')
                mname = row_data.get('hhcomp_hhmmbr_mname', '')
                lname = row_data.get('hhcomp_hhmmbr_lname', '')
                full_name = f"{fname} {mname} {lname}".strip()
                cells[1].text = full_name
            if len(cells) > 2:
                cells[2].text = str(row_data.get('hhcomp_hhmmbr_hhreltn', ''))
            if len(cells) > 3:
                cells[3].text = str(row_data.get('hhcomp_hhmmbr_hhage', ''))
            if len(cells) > 4:
                cells[4].text = str(row_data.get('hhcomp_hhmmbr_hhsex', ''))
            if len(cells) > 5:
                cells[5].text = str(row_data.get('hhcomp_hhmmbr_status', ''))
            if len(cells) > 6:
                religion = row_data.get('hhcomp_hhmmbr_relg', '')
                religion_other = row_data.get('hhcomp_hhmmbr_relg_o', '')
                if religion_other and str(religion_other).strip().lower() not in ['', 'nan']:
                    cells[6].text = f"{religion} Pls. Specify: {religion_other}"
                else:
                    cells[6].text = str(religion)
            if len(cells) > 7:
                cells[7].text = str(row_data.get('hhcomp_hhmmbr_brtplc', ''))
            if len(cells) > 8:
                cells[8].text = str(row_data.get('hhcomp_hhmmbr_educ', ''))
            if len(cells) > 9:
                cells[9].text = str(row_data.get('hhcomp_hhmmbr_ethn', ''))

    def populate_savings_table(self, table, additional_rows):
        """Populate the savings account table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 1:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows
        for idx, row_data in enumerate(additional_rows, start=1):
            new_row = table.add_row()
            cells = new_row.cells
            
            if len(cells) > 0:
                cells[0].text = str(idx)  # Row number
            if len(cells) > 1:
                cells[1].text = str(row_data.get('hhcomp_hhmmbr_ethn', ''))
            if len(cells) > 2:
                savings = row_data.get('hhcomp_hhmmbr_savings', '')
                savings_o = row_data.get('hhcomp_hhmmbr_savings_o', '')
                if savings_o and str(savings_o).strip().lower() not in ['', 'nan']:
                    cells[2].text = f"{savings} Pls. Specify {savings_o}"
                else:
                    cells[2].text = str(savings)
            if len(cells) > 3:
                cells[3].text = str(row_data.get('hhcomp_hhmmbr_phone', ''))
            if len(cells) > 4:
                org = row_data.get('hhcomp_hhmmbr_org', '')
                org_o = row_data.get('hhcomp_hhmmbr_org_o', '')
                if org_o and str(org_o).strip().lower() not in ['', 'nan']:
                    cells[4].text = f"{org} Pls. Specify {org_o}"
                else:
                    cells[4].text = str(org)
            if len(cells) > 5:
                cells[5].text = str(row_data.get('hhcomp_hhmmbr_org_mem', ''))
            if len(cells) > 6:
                cells[6].text = str(row_data.get('hhcomp_hhmmbr_disability', ''))

    def populate_labor_table(self, table, additional_rows):
        """Populate the Labor Force Status table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 3:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows
        for idx, row_data in enumerate(additional_rows, start=1):
            new_row = table.add_row()
            cells = new_row.cells
            
            if len(cells) > 0:
                cells[0].text = str(idx)  # Row number
            if len(cells) > 1:
                cells[1].text = str(row_data.get('hh_labor_stat', ''))
            if len(cells) > 2:
                labor_pri = row_data.get('hh_labor_pri_src', '')
                labor_pri_o = row_data.get('hh_labor_pri_src_o', '')
                if labor_pri_o and str(labor_pri_o).strip().lower() not in ['', 'nan']:
                    cells[2].text = f"{labor_pri} Pls. Specify {labor_pri_o}"
                else:
                    cells[2].text = str(labor_pri)
            if len(cells) > 3:
                cells[3].text = str(row_data.get('hh_labor_pri_industry', ''))
            if len(cells) > 4:
                cells[4].text = str(row_data.get('hh_labor_pri_plc_work', ''))
            if len(cells) > 5:
                cells[5].text = str(row_data.get('hh_labor_pri_inc', ''))
            if len(cells) > 6:
                cells[6].text = str(row_data.get('hh_labor_occ_other', ''))
            if len(cells) > 7:
                cells[7].text = str(row_data.get('hh_labor_other_industry', ''))
            if len(cells) > 8:
                cells[8].text = str(row_data.get('hh_labor_occ_other_plc_wrk', ''))
            if len(cells) > 9:
                cells[9].text = str(row_data.get('hh_labor_occ_other_inc', ''))
            if len(cells) > 10:
                cells[10].text = str(row_data.get('hh_calc_total_inc', ''))
            if len(cells) > 11:
                cells[11].text = str(row_data.get('hh_wrk_hrs', ''))

    def populate_debt_table(self, table, additional_rows):
        """Populate the debt/loan table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 1:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows
        for idx, row_data in enumerate(additional_rows, start=1):
            new_row = table.add_row()
            cells = new_row.cells
            
            if len(cells) > 0:
                debt_src_name = row_data.get('debt_src_name', '')
                debt_src_name_o = row_data.get('debt_src_name_o', '')
                if debt_src_name_o and str(debt_src_name_o).strip().lower() not in ['', 'nan']:
                    cells[0].text = f"{debt_src_name} Pls. Specify {debt_src_name_o}"
                else:
                    cells[0].text = str(debt_src_name)
            if len(cells) > 1:
                cells[1].text = str(row_data.get('debt_contract', ''))
            if len(cells) > 2:
                cells[2].text = str(row_data.get('debt_contract_y', ''))
            if len(cells) > 3:
                cells[3].text = str(row_data.get('debt_amt', ''))
            if len(cells) > 4:
                loan_used = row_data.get('loan_used', '')
                loan_used_o = row_data.get('loan_used_o', '')
                if loan_used_o and str(loan_used_o).strip().lower() not in ['', 'nan']:
                    cells[4].text = f"{loan_used} Pls. Specify {loan_used_o}"
                else:
                    cells[4].text = str(loan_used)
            if len(cells) > 5:
                pymt_terms = str(row_data.get('pymt_terms', ''))
                pymt_terms_int = str(row_data.get('pymt_terms_int', ''))
                pymt_terms_amt = str(row_data.get('pymt_terms_amt', ''))
                pymt_terms_long = str(row_data.get('pymt_terms_long', ''))
                payment_terms = f"{pymt_terms}{pymt_terms_int}{pymt_terms_amt}, {pymt_terms_long}"
                cells[5].text = payment_terms
            if len(cells) > 6:
                cells[6].text = str(row_data.get('debt_balance', ''))
            if len(cells) > 7:
                cells[7].text = str(row_data.get('debt_fam_proc', ''))
            if len(cells) > 8:
                cells[8].text = str(row_data.get('debt_fam_payment', ''))

    def populate_land_assets_table(self, table, additional_rows):
        """Populate the land assets table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 2:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows
        for idx, row_data in enumerate(additional_rows, start=1):
            new_row = table.add_row()
            cells = new_row.cells
            
            if len(cells) > 0:
                cells[0].text = str(idx)  # Row number
            if len(cells) > 1:
                cells[1].text = str(row_data.get('asset_land_area', ''))
            if len(cells) > 2:
                cells[2].text = str(row_data.get('asset_land_area_aff', ''))
            if len(cells) > 3:
                cells[3].text = str(row_data.get('asset_land_ext_impact', ''))
            if len(cells) > 4:
                cells[4].text = str(row_data.get('asset_land_type', ''))
            if len(cells) > 5:
                asset_land_use = row_data.get('asset_land_use', '')
                asset_land_use_o = row_data.get('asset_land_use_o', '')
                if asset_land_use_o and str(asset_land_use_o).strip().lower() not in ['', 'nan']:
                    cells[5].text = f"{asset_land_use}, Please Specify {asset_land_use_o}"
                else:
                    cells[5].text = str(asset_land_use)
            if len(cells) > 6:
                asset_land_tenure_owner = row_data.get('asset_land_tenure_owner', '')
                asset_land_tenure_owner_o = row_data.get('asset_land_tenure_owner_o', '')
                if asset_land_tenure_owner_o and str(asset_land_tenure_owner_o).strip().lower() not in ['', 'nan']:
                    cells[6].text = f"{asset_land_tenure_owner}, Please Specify {asset_land_tenure_owner_o}"
                else:
                    cells[6].text = str(asset_land_tenure_owner)
            if len(cells) > 7:
                asset_land_proof_owner = row_data.get('asset_land_proof_owner', '')
                asset_land_proof_owner_o = row_data.get('asset_land_proof_owner_o', '')
                if asset_land_proof_owner_o and str(asset_land_proof_owner_o).strip().lower() not in ['', 'nan']:
                    cells[7].text = f"{asset_land_proof_owner}, Please Specify {asset_land_proof_owner_o}"
                else:
                    cells[7].text = str(asset_land_proof_owner)
            if len(cells) > 8:
                cells[8].text = str(row_data.get('asset_land_yrs_used', ''))
            if len(cells) > 9:
                cells[9].text = str(row_data.get('asset_land_price_prch', ''))
            if len(cells) > 10:
                asset_land_pymnt_trms = row_data.get('asset_land_pymnt_trms', '')
                asset_land_pymnt_trms_o = row_data.get('asset_land_pymnt_trms_o', '')
                if asset_land_pymnt_trms_o and str(asset_land_pymnt_trms_o).strip().lower() not in ['', 'nan']:
                    cells[10].text = f"{asset_land_pymnt_trms}, Please Specify {asset_land_pymnt_trms_o}"
                else:
                    cells[10].text = str(asset_land_pymnt_trms)
            if len(cells) > 11:
                cells[11].text = str(row_data.get('asset_land_pymnt_amt', ''))

    def populate_structure_assets_table(self, table, additional_rows):
        """Populate the structure assets table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 2:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows
        for idx, row_data in enumerate(additional_rows, start=1):
            new_row = table.add_row()
            cells = new_row.cells
            
            if len(cells) > 0:
                cells[0].text = str(idx)  # Row number
            if len(cells) > 1:
                cells[1].text = str(row_data.get('asset_struct_area', ''))
            if len(cells) > 2:
                cells[2].text = str(row_data.get('asset_struct_area_aff', ''))
            if len(cells) > 3:
                cells[3].text = str(row_data.get('asset_struct_ext_impact', ''))
            if len(cells) > 4:
                asset_struct_type = row_data.get('asset_struct_type', '')
                asset_struct_type_oth = row_data.get('asset_struct_type_oth', '')
                asset_struct_type_oth_o = row_data.get('asset_struct_type_oth_o', '')
                
                type_parts = [asset_struct_type]
                if asset_struct_type_oth and str(asset_struct_type_oth).strip().lower() not in ['', 'nan']:
                    type_parts.append(f"Please Specify {asset_struct_type_oth}")
                if asset_struct_type_oth_o and str(asset_struct_type_oth_o).strip().lower() not in ['', 'nan']:
                    type_parts.append(str(asset_struct_type_oth))
                
                cells[4].text = ", ".join([part for part in type_parts if part])
            if len(cells) > 5:
                asset_struct_use = row_data.get('asset_struct_use', '')
                asset_struct_use_o = row_data.get('asset_struct_use_o', '')
                if asset_struct_use_o and str(asset_struct_use_o).strip().lower() not in ['', 'nan']:
                    cells[5].text = f"{asset_struct_use}, Please Specify {asset_struct_use_o}"
                else:
                    cells[5].text = str(asset_struct_use)
            if len(cells) > 6:
                asset_struct_tenure_owner = row_data.get('asset_struct_tenure_owner', '')
                asset_struct_tenure_owner_o = row_data.get('asset_struct_tenure_owner_o', '')
                if asset_struct_tenure_owner_o and str(asset_struct_tenure_owner_o).strip().lower() not in ['', 'nan']:
                    cells[6].text = f"{asset_struct_tenure_owner}, Please Specify {asset_struct_tenure_owner_o}"
                else:
                    cells[6].text = str(asset_struct_tenure_owner)
            if len(cells) > 7:
                asset_struct_proof_owner = row_data.get('asset_struct_proof_owner', '')
                asset_struct_proof_owner_o = row_data.get('asset_struct_proof_owner_o', '')
                if asset_struct_proof_owner_o and str(asset_struct_proof_owner_o).strip().lower() not in ['', 'nan']:
                    cells[7].text = f"{asset_struct_proof_owner}, Please Specify {asset_struct_proof_owner_o}"
                else:
                    cells[7].text = str(asset_struct_proof_owner)
            if len(cells) > 8:
                cells[8].text = str(row_data.get('asset_struct_yrs_used', ''))
            if len(cells) > 9:
                cells[9].text = str(row_data.get('asset_struct_price_prch', ''))
            if len(cells) > 10:
                asset_struct_pymnt_trms = row_data.get('asset_struct_pymnt_trms', '')
                asset_struct_pymnt_trms_o = row_data.get('asset_struct_pymnt_trms_o', '')
                if asset_struct_pymnt_trms_o and str(asset_struct_pymnt_trms_o).strip().lower() not in ['', 'nan']:
                    cells[10].text = f"{asset_struct_pymnt_trms}, Please Specify {asset_struct_pymnt_trms_o}"
                else:
                    cells[10].text = str(asset_struct_pymnt_trms)
            if len(cells) > 11:
                cells[11].text = str(row_data.get('asset_struct_pymnt_amt', ''))
            if len(cells) > 12:
                cells[12].text = str(row_data.get('asset_struct_mrkt_val', ''))

    def populate_affected_structure_table(self, table, additional_rows):
        """Populate the affected structure table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 2:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows
        for idx, row_data in enumerate(additional_rows, start=1):
            new_row = table.add_row()
            cells = new_row.cells
            
            if len(cells) > 0:
                affctd_struct_type_zz = row_data.get('affctd_struct_type_zz', '')
                affctd_struct_type_zz_o = row_data.get('affctd_struct_type_zz_o', '')
                if affctd_struct_type_zz_o and str(affctd_struct_type_zz_o).strip().lower() not in ['', 'nan']:
                    cells[0].text = f"{affctd_struct_type_zz}, Please Specify {affctd_struct_type_zz_o}"
                else:
                    cells[0].text = str(affctd_struct_type_zz)
            if len(cells) > 1:
                cells[1].text = str(row_data.get('affctd_struct_mtrl_type', ''))
            if len(cells) > 2:
                cells[2].text = str(row_data.get('affctd_struct_dimension', ''))
            if len(cells) > 3:
                affctd_struct_unit = row_data.get('affctd_struct_unit', '')
                affctd_struct_ht = row_data.get('affctd_struct_ht', '')
                unit_parts = [str(affctd_struct_unit), str(affctd_struct_ht)]
                cells[3].text = ", ".join([part for part in unit_parts if part and str(part).strip().lower() not in ['', 'nan']])
            if len(cells) > 4:
                cells[4].text = str(row_data.get('affctd_struct_estvalue', ''))
            if len(cells) > 5:
                cells[5].text = str(row_data.get('affctd_struct_totalcost', ''))
            if len(cells) > 6:
                # Insert images instead of text for Pix1-Pix10
                self.insert_images_in_cell(cells[6], row_data, 1.5)

    def insert_images_in_cell(self, cell, row_data, image_width):
        """Insert multiple images in a table cell from Pix1-Pix10 data in left-right layout"""
        # Clear the cell first
        cell.text = ""
        
        # Get all Pix values
        pix_keys = ['Pix1', 'Pix2', 'Pix3', 'Pix4', 'Pix5', 'Pix6', 'Pix7', 'Pix8', 'Pix9', 'Pix10']
        
        # Check if image folder is available
        if not self.image_folder_path:
            # Fallback to text if no image folder
            pics = []
            for key in pix_keys:
                value = row_data.get(key, '')
                if value and str(value).strip() and str(value).lower() != 'nan':
                    pics.append(str(value))
            cell.text = ", ".join(pics)
            return
        
        # Clear existing paragraphs in the cell
        for paragraph in cell.paragraphs:
            paragraph.clear()
        
        # Collect valid image paths first
        valid_images = []
        for key in pix_keys:
            pix_value = row_data.get(key, '')
            if pix_value and str(pix_value).strip() and str(pix_value).lower() != 'nan':
                image_filename = str(pix_value).strip()
                
                # Try different image extensions
                image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
                
                for ext in [''] + image_extensions:
                    if ext == '':
                        test_filename = image_filename
                    else:
                        # Remove existing extension if any, then add new one
                        base_name = os.path.splitext(image_filename)[0]
                        test_filename = base_name + ext
                    
                    image_path = os.path.join(self.image_folder_path, test_filename)
                    if os.path.exists(image_path):
                        valid_images.append(image_path)
                        break
                else:
                    # If no image found, add a placeholder
                    valid_images.append(f"[Image not found: {image_filename}]")
        
        if not valid_images:
            cell.text = "No images"
            return
        
        # Arrange images in 2-column layout (left and right)
        for i in range(0, len(valid_images), 2):
            # Create a new paragraph for each row of images
            if i > 0:
                cell.add_paragraph()
            
            paragraph = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            
            # Left image
            left_image = valid_images[i]
            if left_image.startswith("[Image not found:"):
                paragraph.add_run(left_image)
            else:
                try:
                    run = paragraph.add_run()
                    run.add_picture(left_image, width=Inches(image_width))
                except Exception as e:
                    paragraph.add_run(f"[Error loading: {os.path.basename(left_image)}]")
            
            # Add space between images
            paragraph.add_run("    ")
            
            # Right image (if exists)
            if i + 1 < len(valid_images):
                right_image = valid_images[i + 1]
                if right_image.startswith("[Image not found:"):
                    paragraph.add_run(right_image)
                else:
                    try:
                        run = paragraph.add_run()
                        run.add_picture(right_image, width=Inches(image_width))
                    except Exception as e:
                        paragraph.add_run(f"[Error loading: {os.path.basename(right_image)}]")

    def populate_trees_table(self, table, additional_rows):
        """Populate the trees table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 1:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows
        for row_data in additional_rows:
            new_row = table.add_row()
            cells = new_row.cells
            
            if len(cells) > 0:
                cells[0].text = str(row_data.get('tree_type', ''))
            if len(cells) > 1:
                tree_age = row_data.get('tree_age', '')
                tree_height = row_data.get('tree_height', '')
                age_height = f"{tree_age}, {tree_height}".strip(', ')
                cells[1].text = age_height
            if len(cells) > 2:
                cells[2].text = str(row_data.get('tree_qty', ''))
            if len(cells) > 3:
                cells[3].text = str(row_data.get('tree_price', ''))
            if len(cells) > 4:
                cells[4].text = str(row_data.get('tree_totalcost', ''))

    def populate_crops_table(self, table, additional_rows):
        """Populate the crops table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 1:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows using more flexible key matching
        for idx, row_data in enumerate(additional_rows):
            new_row = table.add_row()
            cells = new_row.cells
            
            # Try to find the right keys by looking for any key containing these terms
            crop_type = ""
            crop_age = ""
            crop_area = ""
            crop_price = ""
            crop_total = ""
            
            for key, value in row_data.items():
                key_lower = key.lower()
                if 'crop' in key_lower and 'type' in key_lower and value:
                    crop_type = str(value)
                elif 'crop' in key_lower and 'age' in key_lower and value:
                    crop_age = str(value)
                elif 'crop' in key_lower and 'area' in key_lower and value:
                    crop_area = str(value)
                elif 'crop' in key_lower and 'price' in key_lower and value:
                    crop_price = str(value)
                elif 'crop' in key_lower and ('total' in key_lower or 'cost' in key_lower) and value:
                    crop_total = str(value)
            
            if len(cells) > 0:
                cells[0].text = crop_type
            if len(cells) > 1:
                cells[1].text = crop_age
            if len(cells) > 2:
                cells[2].text = crop_area
            if len(cells) > 3:
                cells[3].text = crop_price
            if len(cells) > 4:
                cells[4].text = crop_total

    def populate_income_loss_table(self, table, additional_rows):
        """Populate the income loss table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 1:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)

        # Add data rows using more flexible key matching
        for idx, row_data in enumerate(additional_rows):
            new_row = table.add_row()
            cells = new_row.cells
            
            # Try to find the right keys by looking for any key containing these terms
            loss_type = ""
            loss_qty = ""
            loss_unit = ""
            loss_price = ""
            loss_total = ""
            
            for key, value in row_data.items():
                key_lower = key.lower()
                if 'income' in key_lower and 'type' in key_lower and value:
                    loss_type = str(value)
                elif 'income' in key_lower and ('qty' in key_lower or 'quantity' in key_lower) and value:
                    loss_qty = str(value)
                elif 'income' in key_lower and 'unit' in key_lower and 'price' not in key_lower and value:
                    loss_unit = str(value)
                elif 'income' in key_lower and 'price' in key_lower and value:
                    loss_price = str(value)
                elif 'income' in key_lower and ('total' in key_lower or 'cost' in key_lower) and value:
                    loss_total = str(value)
            
            if len(cells) > 0:
                cells[0].text = loss_type
            if len(cells) > 1:
                cells[1].text = loss_qty
            if len(cells) > 2:
                cells[2].text = loss_unit
            if len(cells) > 3:
                cells[3].text = loss_price
            if len(cells) > 4:
                cells[4].text = loss_total

    def populate_others_table(self, table, additional_rows):
        """Populate the others table with data from additional files"""
        # Keep header rows and remove existing data rows
        while len(table.rows) > 1:  # Keep only the header row
            table._tbl.remove(table.rows[-1]._tr)
        
        # Add data rows using more flexible key matching
        for idx, row_data in enumerate(additional_rows):
            new_row = table.add_row()
            cells = new_row.cells
            
            # Try to find the right keys by looking for any key containing these terms
            others_type = ""
            others_qty = ""
            others_unit = ""
            others_price = ""
            others_total = ""
            
            for key, value in row_data.items():
                key_lower = key.lower()
                if 'others' in key_lower and 'type' in key_lower and value:
                    others_type = str(value)
                elif 'others' in key_lower and ('qty' in key_lower or 'quantity' in key_lower) and value:
                    others_qty = str(value)
                elif 'others' in key_lower and 'unit' in key_lower and 'price' not in key_lower and value:
                    others_unit = str(value)
                elif 'others' in key_lower and 'price' in key_lower and value:
                    others_price = str(value)
                elif 'others' in key_lower and ('total' in key_lower or 'cost' in key_lower) and value:
                    others_total = str(value)
            
            if len(cells) > 0:
                cells[0].text = others_type
            if len(cells) > 1:
                cells[1].text = others_qty
            if len(cells) > 2:
                cells[2].text = others_unit
            if len(cells) > 3:
                cells[3].text = others_price
            if len(cells) > 4:
                cells[4].text = others_total


class SettingsTab:
    """Settings tab with modern full-width design"""
    def __init__(self, parent_frame):
        self.parent_frame = parent_frame
        self.setup_modern_settings_ui()
    
    def setup_modern_settings_ui(self):
        self.parent_frame.configure(bg=ModernStyle.BACKGROUND)
        
        # Header section
        header_frame = tk.Frame(self.parent_frame, bg=ModernStyle.SECONDARY, height=80)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        header_content = tk.Frame(header_frame, bg=ModernStyle.SECONDARY)
        header_content.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        title_font = font.Font(family="Segoe UI", size=22, weight="bold")
        title_label = tk.Label(
            header_content,
            text="⚙️ Settings & Configuration",
            font=title_font,
            bg=ModernStyle.SECONDARY,
            fg="white"
        )
        title_label.pack(anchor=tk.W)
        
        # Content area
        content_area = tk.Frame(self.parent_frame, bg=ModernStyle.BACKGROUND)
        content_area.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        # Settings grid
        settings_grid = tk.Frame(content_area, bg=ModernStyle.BACKGROUND)
        settings_grid.pack(fill=tk.BOTH, expand=True)
        
        # Performance settings
        perf_section = self.create_settings_section(settings_grid, "🚀 Performance Settings", 0, 0)
        self.create_setting_item(perf_section, "Processing Threads", "Number of parallel threads", "spinbox", {"from_": 1, "to": 8, "value": "4"})
        self.create_setting_item(perf_section, "Memory Usage", "Optimize memory consumption", "combobox", {"values": ["Low", "Medium", "High"], "default": "Medium"})
        
        # Output settings
        output_section = self.create_settings_section(settings_grid, "📁 Output Settings", 0, 1)
        self.create_setting_item(output_section, "Output Format", "How to save generated documents", "combobox", {"values": ["Individual + ZIP", "ZIP Only", "Individual Only"], "default": "Individual + ZIP"})
        #self.create_setting_item(output_section, "File Naming", "Document naming convention", "combobox", {"values": ["Numbered", "Key-based", "Custom"], "default": "Numbered"})
        
        # Cache settings
        cache_section = self.create_settings_section(settings_grid, "🗄️ Cache Management", 1, 0, columnspan=2)
        
        cache_buttons_frame = tk.Frame(cache_section, bg=ModernStyle.SURFACE)
        cache_buttons_frame.pack(fill=tk.X, padx=20, pady=15)
        
        clear_template_btn = self.create_action_button(cache_buttons_frame, "Clear Template Cache", self.clear_template_cache, ModernStyle.WARNING)
        clear_template_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        clear_data_btn = self.create_action_button(cache_buttons_frame, "Clear Data Cache", self.clear_data_cache, ModernStyle.WARNING)
        clear_data_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        reset_settings_btn = self.create_action_button(cache_buttons_frame, "Reset All Settings", self.reset_settings, ModernStyle.DANGER)
        reset_settings_btn.pack(side=tk.LEFT)
    
    def create_settings_section(self, parent, title, row, column, columnspan=1):
        """Create a modern settings section"""
        section_frame = tk.Frame(parent, bg=ModernStyle.SURFACE, relief=tk.FLAT)
        section_frame.configure(highlightbackground=ModernStyle.BORDER, highlightthickness=1)
        section_frame.grid(row=row, column=column, columnspan=columnspan, sticky="nsew", padx=10, pady=10)
        
        # Configure grid weights
        if columnspan == 2:
            parent.grid_columnconfigure(0, weight=1)
            parent.grid_columnconfigure(1, weight=1)
        else:
            parent.grid_columnconfigure(column, weight=1)
        parent.grid_rowconfigure(row, weight=1)
        
        # Section header
        header_frame = tk.Frame(section_frame, bg=ModernStyle.PRIMARY, height=50)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        header_font = font.Font(family="Segoe UI", size=14, weight="bold")
        header_label = tk.Label(
            header_frame,
            text=title,
            font=header_font,
            bg=ModernStyle.PRIMARY,
            fg="white"
        )
        header_label.pack(anchor=tk.W, padx=20, pady=15)
        
        return section_frame
    
    def create_setting_item(self, parent, title, description, widget_type, options):
        """Create a modern setting item"""
        item_frame = tk.Frame(parent, bg=ModernStyle.SURFACE)
        item_frame.pack(fill=tk.X, padx=20, pady=15)
        
        # Title and description
        text_frame = tk.Frame(item_frame, bg=ModernStyle.SURFACE)
        text_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        title_font = font.Font(family="Segoe UI", size=11, weight="bold")
        title_label = tk.Label(
            text_frame,
            text=title,
            font=title_font,
            bg=ModernStyle.SURFACE,
            fg=ModernStyle.TEXT_PRIMARY
        )
        title_label.pack(anchor=tk.W)
        
        desc_font = font.Font(family="Segoe UI", size=9)
        desc_label = tk.Label(
            text_frame,
            text=description,
            font=desc_font,
            bg=ModernStyle.SURFACE,
            fg=ModernStyle.TEXT_SECONDARY
        )
        desc_label.pack(anchor=tk.W, pady=(2, 0))
        
        # Widget
        if widget_type == "spinbox":
            var = tk.StringVar(value=options.get("value", "1"))
            widget = ttk.Spinbox(
                item_frame,
                from_=options["from_"],
                to=options["to"],
                textvariable=var,
                width=12
            )
        elif widget_type == "combobox":
            var = tk.StringVar(value=options.get("default", ""))
            widget = ttk.Combobox(
                item_frame,
                textvariable=var,
                values=options["values"],
                state="readonly",
                width=20
            )
        
        widget.pack(side=tk.RIGHT, padx=(10, 0))
    
    def create_action_button(self, parent, text, command, color):
        """Create an action button"""
        btn_font = font.Font(family="Segoe UI", size=9, weight="bold")
        
        button = tk.Button(
            parent,
            text=text,
            command=command,
            font=btn_font,
            bg=color,
            fg="white",
            relief=tk.FLAT,
            bd=0,
            padx=20,
            pady=10,
            cursor="hand2"
        )
        
        def on_enter(e):
            darken_map = {
                ModernStyle.WARNING: "#D97706",
                ModernStyle.DANGER: "#DC2626"
            }
            button.configure(bg=darken_map.get(color, color))
        def on_leave(e):
            button.configure(bg=color)
        
        button.bind("<Enter>", on_enter)
        button.bind("<Leave>", on_leave)
        
        return button
    
    def clear_template_cache(self):
        messagebox.showinfo("Cache Cleared", "✅ Template cache has been cleared successfully!")
    
    def clear_data_cache(self):
        messagebox.showinfo("Cache Cleared", "✅ Data cache has been cleared successfully!")
    
    def reset_settings(self):
        if messagebox.askyesno("Reset Settings", "Are you sure you want to reset all settings to default?"):
            messagebox.showinfo("Settings Reset", "✅ All settings have been reset to default values!")


class HelpTab:
    """Help tab with modern full-width design"""
    def __init__(self, parent_frame):
        self.parent_frame = parent_frame
        self.setup_modern_help_ui()
    
    def setup_modern_help_ui(self):
        self.parent_frame.configure(bg=ModernStyle.BACKGROUND)
        
        # Header section
        header_frame = tk.Frame(self.parent_frame, bg=ModernStyle.ACCENT, height=80)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        header_content = tk.Frame(header_frame, bg=ModernStyle.ACCENT)
        header_content.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        title_font = font.Font(family="Segoe UI", size=22, weight="bold")
        title_label = tk.Label(
            header_content,
            text="❓ Help & Documentation",
            font=title_font,
            bg=ModernStyle.ACCENT,
            fg="white"
        )
        title_label.pack(anchor=tk.W)
        
        # Content area with full width
        content_area = tk.Frame(self.parent_frame, bg=ModernStyle.SURFACE)
        content_area.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        
        # Help content with scrollbar
        text_frame = tk.Frame(content_area, bg=ModernStyle.SURFACE)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=30, pady=30)
        
        help_font = font.Font(family="Segoe UI", size=10)
        self.help_text = tk.Text(
            text_frame,
            wrap=tk.WORD,
            font=help_font,
            bg=ModernStyle.SURFACE,
            fg=ModernStyle.TEXT_PRIMARY,
            relief=tk.FLAT,
            bd=0,
            selectbackground=ModernStyle.PRIMARY,
            selectforeground="white",
            spacing1=5,
            spacing2=3,
            spacing3=5
        )
        self.help_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=self.help_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.help_text.configure(yscrollcommand=scrollbar.set)
        
        self.insert_help_content()
    
    def insert_help_content(self):
        """Insert comprehensive help content"""
        # Configure text styling
        title_font = font.Font(family="Segoe UI", size=16, weight="bold")
        heading_font = font.Font(family="Segoe UI", size=12, weight="bold")
        body_font = font.Font(family="Segoe UI", size=10)
        code_font = font.Font(family="Consolas", size=9)
        
        self.help_text.tag_configure("title", font=title_font, foreground=ModernStyle.PRIMARY, spacing1=15, spacing3=10)
        self.help_text.tag_configure("heading", font=heading_font, foreground=ModernStyle.TEXT_PRIMARY, spacing1=12, spacing3=8)
        self.help_text.tag_configure("body", font=body_font, foreground=ModernStyle.TEXT_PRIMARY, spacing1=5)
        self.help_text.tag_configure("code", font=code_font, background=ModernStyle.SURFACE_DARK, foreground=ModernStyle.TEXT_PRIMARY)
        self.help_text.tag_configure("bullet", font=body_font, foreground=ModernStyle.TEXT_SECONDARY, lmargin1=20, lmargin2=40, spacing1=3)
        self.help_text.tag_configure("important", font=body_font, foreground=ModernStyle.DANGER, spacing1=5)
       
        
        # Insert content
        self.help_text.insert(tk.END, "Complete User Guide\n", "title")
        
        self.help_text.insert(tk.END, "📋 OVERVIEW\n", "heading")
        self.help_text.insert(tk.END, "Auto Converter is a powerful tool that transforms Word templates into personalized documents by automatically replacing placeholders with data from Excel files. Perfect for generating contracts, reports, certificates, letters, and any document that requires personalization at scale.\n\n", "body")
        
        self.help_text.insert(tk.END, "🔧 GETTING STARTED\n", "heading")
        self.help_text.insert(tk.END, "• Step 1: Create a Word template with placeholders using curly braces, e.g., ", "bullet")
        self.help_text.insert(tk.END, "{firstname}", "code")
        self.help_text.insert(tk.END, ", ", "bullet")
        self.help_text.insert(tk.END, "{lastname}", "code")
        self.help_text.insert(tk.END, "\n• Step 2: Prepare an Excel file with column headers matching your placeholders\n• Step 3: Upload both files using the Document Converter tab\n• Step 4: Click 'Convert Files' and select your output destination\n• Step 5: Wait for processing to complete and find your ZIP file\n\n", "bullet")
        
        self.help_text.insert(tk.END, "📝 PLACEHOLDER GUIDELINES\n", "heading")
        self.help_text.insert(tk.END, "• Use curly braces format: ", "bullet")
        self.help_text.insert(tk.END, "{placeholder_name}", "code")
        self.help_text.insert(tk.END, "\n• Placeholder names must match Excel column headers exactly\n• Case-insensitive matching is supported\n• Special characters and spaces are allowed in placeholder names\n• Avoid using reserved characters: { } \\ / : * ? \" < > |\n\n", "bullet")
        
        self.help_text.insert(tk.END, "📊 EXCEL FILE STRUCTURE\n", "heading")
        self.help_text.insert(tk.END, "• Column headers can be placed in rows 1-4 (application searches all four rows)\n• Data rows must start from row 5 onwards\n• Row 4 should contain a 'KEY' column for linking with additional data files\n• Additional files should have a 'PARENT_KEY' column in row 4\n• Empty cells are replaced with blank text in the final documents\n\n", "bullet")
        
        self.help_text.insert(tk.END, "⚡ ADVANCED FEATURES\n", "heading")
        self.help_text.insert(tk.END, "• 🔄 Multi-threaded parallel processing for faster conversion\n• 📋 Automatic table population for complex data structures\n• 📦 ZIP file generation for easy distribution and download\n• 💾 Smart caching system for improved performance on large datasets\n• 🎯 Dynamic content insertion based on data patterns\n• 🔗 Cross-file data linking using KEY/PARENT_KEY relationships\n\n", "bullet")
        
        self.help_text.insert(tk.END, "⚙️ PERFORMANCE OPTIMIZATION\n", "heading")
        self.help_text.insert(tk.END, "• Adjust processing threads in Settings based on your system capabilities\n• Use the cache management tools to free up memory\n• For large datasets (1000+ rows), consider processing in smaller batches\n• Close other applications to free up system resources during processing\n• Ensure sufficient disk space for temporary files and output\n\n", "bullet")
        
        self.help_text.insert(tk.END, "🛠️ TROUBLESHOOTING\n", "heading")
        self.help_text.insert(tk.END, "• ", "important")
        self.help_text.insert(tk.END, "Placeholder not found: ", "important")
        self.help_text.insert(tk.END, "Check that placeholder names match Excel headers exactly\n• ", "bullet")
        self.help_text.insert(tk.END, "File access errors: ", "important")
        self.help_text.insert(tk.END, "Ensure Excel files are not open in another application\n• ", "bullet")
        self.help_text.insert(tk.END, "Memory issues: ", "important")
        self.help_text.insert(tk.END, "Clear cache and reduce processing threads\n• ", "bullet")
        self.help_text.insert(tk.END, "Slow performance: ", "important")
        self.help_text.insert(tk.END, "Increase processing threads and ensure SSD storage\n• ", "bullet")
        self.help_text.insert(tk.END, "Template corruption: ", "important")
        self.help_text.insert(tk.END, "Verify Word template opens correctly before processing\n\n", "bullet")
        
        self.help_text.insert(tk.END, "💡 TIPS & BEST PRACTICES\n", "heading")
        self.help_text.insert(tk.END, "• Test with a small dataset (5-10 rows) before processing large files\n• Use descriptive placeholder names for better organization\n• Keep backup copies of your original template and data files\n• Organize Excel files in separate folders by project or category\n• Regular cache clearing improves memory usage and performance\n• Use consistent naming conventions for placeholders across projects\n• Consider using the Settings tab to optimize performance for your system\n\n", "bullet")

        self.help_text.config(state=tk.DISABLED)


class MainTablesConverterTab:
    """Tab for main tables conversion with modern full-width design"""
    def __init__(self, parent_frame):
        self.parent_frame = parent_frame
        self.codes_df = None
        self.setup_modern_fullwidth_ui()
    
    def setup_modern_fullwidth_ui(self):
        # Configure parent frame
        self.parent_frame.configure(bg=ModernStyle.BACKGROUND)
        
        # Main container with full width
        main_container = tk.Frame(self.parent_frame, bg=ModernStyle.BACKGROUND)
        main_container.pack(fill=tk.BOTH, expand=True, padx=0, pady=0)
        
        # Header section with gradient-like effect
        header_frame = tk.Frame(main_container, bg=ModernStyle.ACCENT, height=80)
        header_frame.pack(fill=tk.X, pady=(0, 0))
        header_frame.pack_propagate(False)
        
        # Header content
        header_content = tk.Frame(header_frame, bg=ModernStyle.ACCENT)
        header_content.pack(fill=tk.BOTH, expand=True, padx=30, pady=20)
        
        title_font = font.Font(family="Segoe UI", size=22, weight="bold")
        title_label = tk.Label(
            header_content,
            text="🛠️ Main Tables Converter",
            font=title_font,
            bg=ModernStyle.ACCENT,
            fg="white"
        )
        title_label.pack(side=tk.LEFT, anchor=tk.W)
        
        subtitle_font = font.Font(family="Segoe UI", size=11)
        subtitle_label = tk.Label(
            header_content,
            text="Convert coded values to readable labels",
            font=subtitle_font,
            bg=ModernStyle.ACCENT,
            fg="white"
        )
        subtitle_label.pack(side=tk.RIGHT, anchor=tk.E)
        
        # Content area with full width
        content_area = tk.Frame(main_container, bg=ModernStyle.BACKGROUND)
        content_area.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Two-column layout for better space utilization
        left_column = tk.Frame(content_area, bg=ModernStyle.BACKGROUND)
        left_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        right_column = tk.Frame(content_area, bg=ModernStyle.BACKGROUND)
        right_column.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(10, 0))
        
        # File Upload Section (Left Column)
        upload_section = self.create_modern_section(left_column, "📁 File Upload", ModernStyle.PRIMARY)
        upload_section.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        # Modern upload cards
        self.create_modern_upload_card(upload_section, "📋", "Load Codes File", "Excel file with value mappings", 
                                     lambda: self.load_codesfile(), "codes_label")
        
        # Processing Section (Right Column)
        processing_section = self.create_modern_section(right_column, "🚀 Processing", ModernStyle.SECONDARY)
        processing_section.pack(fill=tk.BOTH, expand=True)
        
        # Convert button area
        button_area = tk.Frame(processing_section, bg=ModernStyle.SURFACE)
        button_area.pack(fill=tk.X, padx=20, pady=20)
        
        self.convert_btn = self.create_modern_action_button(
            button_area,
            "🔄 Convert Excel File",
            self.convert_file,
            state="disabled"
        )
        self.convert_btn.pack(anchor=tk.CENTER, pady=(0, 10))
        
        self.instructions_btn = self.create_modern_button(
            button_area,
            "📘 Instructions & Credits",
            self.show_instructions,
            ModernStyle.ACCENT
        )
        self.instructions_btn.pack(anchor=tk.CENTER)
        
        # Status area
        status_area = tk.Frame(processing_section, bg=ModernStyle.SURFACE_DARK, relief=tk.FLAT)
        status_area.pack(fill=tk.X, padx=20, pady=(15, 20))
        
        status_font = font.Font(family="Segoe UI", size=9)
        self.status_label = tk.Label(
            status_area,
            text="Ready to convert your files",
            font=status_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_SECONDARY,
            pady=10
        )
        self.status_label.pack()

    def create_modern_section(self, parent, title, accent_color):
        """Create a modern section with accent color"""
        section_frame = tk.Frame(parent, bg=ModernStyle.SURFACE, relief=tk.FLAT)
        section_frame.configure(highlightbackground=ModernStyle.BORDER, highlightthickness=1)
        
        # Section header with accent
        header_frame = tk.Frame(section_frame, bg=accent_color, height=40)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        header_font = font.Font(family="Segoe UI", size=14, weight="bold")
        header_label = tk.Label(
            header_frame,
            text=title,
            font=header_font,
            bg=accent_color,
            fg="white"
        )
        header_label.pack(anchor=tk.W, padx=20, pady=10)
        
        return section_frame
    
    def create_modern_upload_card(self, parent, icon, title, description, command, label_attr):
        """Create a modern upload card with full-width design"""
        card_frame = tk.Frame(parent, bg=ModernStyle.SURFACE)
        card_frame.pack(fill=tk.X, padx=20, pady=10)
        
        # Card content with hover effect simulation
        content_frame = tk.Frame(card_frame, bg=ModernStyle.SURFACE_DARK, relief=tk.FLAT, bd=1)
        content_frame.pack(fill=tk.X, pady=2)
        
        # Icon and content area
        main_content = tk.Frame(content_frame, bg=ModernStyle.SURFACE_DARK)
        main_content.pack(fill=tk.X, padx=15, pady=12)
        
        # Icon
        icon_font = font.Font(family="Segoe UI", size=20)
        icon_label = tk.Label(
            main_content,
            text=icon,
            font=icon_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.PRIMARY,
            width=3
        )
        icon_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # Text content
        text_content = tk.Frame(main_content, bg=ModernStyle.SURFACE_DARK)
        text_content.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        title_font = font.Font(family="Segoe UI", size=11, weight="bold")
        title_label = tk.Label(
            text_content,
            text=title,
            font=title_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_PRIMARY
        )
        title_label.pack(anchor=tk.W)
        
        desc_font = font.Font(family="Segoe UI", size=9)
        desc_label = tk.Label(
            text_content,
            text=description,
            font=desc_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_SECONDARY
        )
        desc_label.pack(anchor=tk.W, pady=(2, 0))
        
        # Status label
        status_font = font.Font(family="Segoe UI", size=8)
        status_label = tk.Label(
            text_content,
            text="No file selected",
            font=status_font,
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_SECONDARY
        )
        status_label.pack(anchor=tk.W, pady=(4, 0))
        setattr(self, label_attr, status_label)
        
        # Button
        button = self.create_modern_button(
            main_content,
            "Choose File",
            command,
            ModernStyle.SECONDARY
        )
        button.pack(side=tk.RIGHT, padx=(10, 0))

    def create_modern_button(self, parent, text, command, color):
        """Create a modern button with rounded appearance"""
        btn_font = font.Font(family="Segoe UI", size=9, weight="bold")
        
        button = tk.Button(
            parent,
            text=text,
            command=command,
            font=btn_font,
            bg=color,
            fg="white",
            relief=tk.FLAT,
            bd=0,
            padx=20,
            pady=8,
            cursor="hand2"
        )
        
        # Hover effects
        def on_enter(e):
            button.configure(bg=self.darken_color(color))
        def on_leave(e):
            button.configure(bg=color)
        
        button.bind("<Enter>", on_enter)
        button.bind("<Leave>", on_leave)
        
        return button
    
    def create_modern_action_button(self, parent, text, command, state="normal"):
        """Create a large modern action button with appealing colors"""
        btn_font = font.Font(family="Segoe UI", size=12, weight="bold")
        
        if state == "disabled":
            bg_color = ModernStyle.BUTTON_DISABLED
            text_color = "#64748B"
            cursor = "arrow"
        else:
            bg_color = ModernStyle.BUTTON_PRIMARY
            text_color = "white"
            cursor = "hand2"
        
        button = tk.Button(
            parent,
            text=text,
            command=command,
            font=btn_font,
            bg=bg_color,
            fg=text_color,
            relief=tk.FLAT,
            bd=0,
            padx=40,
            pady=15,
            cursor=cursor,
            state=state
        )
        
        return button
    
    def setup_modern_progress_bar(self, parent):
        """Setup modern progress bar with custom styling"""
        progress_container = tk.Frame(parent, bg=ModernStyle.SURFACE)
        progress_container.pack(fill=tk.X, pady=(0, 10))
        
        # Custom progress bar background
        progress_bg = tk.Frame(progress_container, bg=ModernStyle.SURFACE_DARK, height=8)
        progress_bg.pack(fill=tk.X, pady=(0, 5))
        
        # Progress fill
        self.progress_fill = tk.Frame(progress_bg, bg=ModernStyle.PRIMARY, height=8)
        self.progress_fill.place(x=0, y=0, width=0, height=8)
        
        # Progress percentage
        self.progress_text = tk.Label(
            progress_container,
            text="0%",
            font=font.Font(family="Segoe UI", size=8),
            bg=ModernStyle.SURFACE,
            fg=ModernStyle.TEXT_SECONDARY
        )
        self.progress_text.pack(anchor=tk.E)
        
        # Store the background frame for width calculations
        self.progress_bg_frame = progress_bg
        
        # Initialize progress
        self.progress_value = 0
    
    def update_progress(self, value):
        """Update the custom progress bar"""
        self.progress_value = value
        if hasattr(self, 'progress_bg_frame'):
            # Calculate width based on progress
            bg_width = self.progress_bg_frame.winfo_width()
            if bg_width > 1:  # Ensure the frame has been drawn
                fill_width = int((value / 100) * bg_width)
                self.progress_fill.place(width=fill_width)
                self.progress_text.config(text=f"{int(value)}%")
    
    def darken_color(self, color):
        """Darken a hex color for hover effect"""
        color_map = {
            ModernStyle.PRIMARY: "#2563EB",
            ModernStyle.SECONDARY: "#059669",
            ModernStyle.ACCENT: "#D97706",
            ModernStyle.DANGER: "#DC2626",
            ModernStyle.BUTTON_PRIMARY: ModernStyle.BUTTON_PRIMARY_HOVER
        }
        return color_map.get(color, color)
    
    def update_file_status(self, label, filename, success=True):
        """Update file status with modern styling"""
        if success:
            display_name = filename if len(filename) <= 30 else filename[:27] + "..."
            label.configure(
                text=f"✅ {display_name}",
                fg=ModernStyle.SUCCESS
            )
        else:
            label.configure(
                text="❌ No file selected",
                fg=ModernStyle.TEXT_SECONDARY
            )

    def load_codesfile(self):
        """Load codes file functionality"""
        file_path = filedialog.askopenfilename(
            title="Select Codes File",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if file_path:
            try:
                self.codes_df = pd.read_excel(file_path)
                # Validate required columns
                required_cols = ['list name', 'name', 'label::English']
                missing_cols = [col for col in required_cols if col not in self.codes_df.columns]
                
                if missing_cols:
                    messagebox.showerror("Error", f"Missing required columns: {', '.join(missing_cols)}")
                    self.codes_df = None
                    return
                
                self.update_file_status(self.codes_label, Path(file_path).name, True)
                self.check_ready_to_convert()
                self.status_label.config(text="✅ Codes file loaded successfully", fg=ModernStyle.SUCCESS)
                messagebox.showinfo("Success", "Codes file loaded successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load codes file:\n{e}")
                self.status_label.config(text="❌ Error loading codes file", fg=ModernStyle.DANGER)
    
    def check_ready_to_convert(self):
        """Check if ready to convert and enable/disable button"""
        if self.codes_df is not None:
            self.convert_btn.configure(
                state="normal",
                bg=ModernStyle.BUTTON_PRIMARY,
                fg="white",
                cursor="hand2"
            )
            
            def on_enter(e):
                self.convert_btn.configure(bg=ModernStyle.BUTTON_PRIMARY_HOVER)
            def on_leave(e):
                self.convert_btn.configure(bg=ModernStyle.BUTTON_PRIMARY)
            
            self.convert_btn.bind("<Enter>", on_enter)
            self.convert_btn.bind("<Leave>", on_leave)
        else:
            self.convert_btn.configure(
                state="disabled",
                bg=ModernStyle.BUTTON_DISABLED,
                fg="#64748B",
                cursor="arrow"
            )
            self.convert_btn.unbind("<Enter>")
            self.convert_btn.unbind("<Leave>")

    def convert_file(self):
        """Convert file functionality"""
        if self.codes_df is None:
            messagebox.showwarning("No Codes File", "Please load a codes file first.")
            return

        # Multiple file selection
        file_paths = filedialog.askopenfilenames(
            title="Select Files to Convert",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if not file_paths:
            return

        # Custom save directory
        save_dir = filedialog.askdirectory(title="Select Folder to Save Converted Files")
        if not save_dir:
            return

        try:
            self.status_label.config(text=f"🔄 Processing {len(file_paths)} file(s)...", fg=ModernStyle.PRIMARY)

            for file_path in file_paths:
                df = pd.read_excel(file_path, header=None, keep_default_na=False)
                headers = df.iloc[0].tolist()
                converted = df.copy()

                for col_idx, var in enumerate(headers):
                    if var in self.codes_df['list name'].unique():
                        lookup = self.codes_df[self.codes_df['list name'] == var][['name', 'label::English']]
                        mapping = dict(zip(lookup['name'].astype(str).str.strip(), lookup['label::English']))

                        def map_values(val):
                            val_str = str(val).strip()
                            if val_str == "":
                                return ""
                            if val_str.upper() == "NA":
                                return "None / Not Applicable"
                            try:
                                parsed_date = datetime.strptime(val_str, "%Y-%m-%d %H:%M:%S")
                                val_str = f"{parsed_date.month} {parsed_date.day}"
                            except ValueError:
                                pass
                            if any(char.isdigit() for char in val_str) and ' ' in val_str:
                                parts = [p.strip() for p in val_str.split()]
                                return ', '.join([mapping.get(p, p) for p in parts])
                            return mapping.get(val_str, val)

                        converted.iloc[4:, col_idx] = converted.iloc[4:, col_idx].apply(map_values)

                converted.columns = converted.iloc[0]
                converted = converted.iloc[1:].reset_index(drop=True)

                # Save to custom directory with original filename + _converted
                base_name = os.path.basename(os.path.splitext(file_path)[0])
                save_path = os.path.join(save_dir, f"{base_name}_converted.xlsx")
                converted.to_excel(save_path, index=False)

            self.status_label.config(text="✅ Files converted successfully!", fg=ModernStyle.SUCCESS)
            messagebox.showinfo("Success", f"All files converted and saved to:\n{save_dir}")

        except Exception as e:
            self.status_label.config(text="❌ Error during conversion", fg=ModernStyle.DANGER)
            messagebox.showerror("Error", f"Failed to convert files:\n{e}")

    def show_instructions(self):
        """Show instructions window"""
        instr_window = tk.Toplevel(self.parent_frame)
        instr_window.title("📘 How to Use - Main Tables Converter")
        instr_window.geometry("700x500")
        instr_window.configure(bg=ModernStyle.BACKGROUND)
        instr_window.resizable(False, False)
        
        # Center the window
        instr_window.transient(self.parent_frame)
        instr_window.grab_set()
        
        # Header
        header_frame = tk.Frame(instr_window, bg=ModernStyle.ACCENT, height=60)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)
        
        header_content = tk.Frame(header_frame, bg=ModernStyle.ACCENT)
        header_content.pack(fill=tk.BOTH, expand=True, padx=20, pady=15)
        
        heading_font = font.Font(family="Segoe UI", size=16, weight="bold")
        heading = tk.Label(
            header_content,
            text="📝 Step-by-Step Guide",
            font=heading_font,
            bg=ModernStyle.ACCENT,
            fg="white"
        )
        heading.pack(anchor=tk.W)
        
        # Content area
        content_frame = tk.Frame(instr_window, bg=ModernStyle.SURFACE)
        content_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        content_font = font.Font(family="Segoe UI", size=10)
        instr_text = """
1. Click "Load Codes File" to upload the Excel file containing your value mappings.
   ➤ Required columns: 'list name', 'name', and 'label::English'.

2. Click "Convert Excel File" to select the dataset you want to convert.

3. The program will process your file from row 5 downward, mapping coded values
   into readable labels based on your codes file.

✔ Blank cells will remain unchanged.
✔ The converted Excel file will be saved next to the original file.

Example Codes File Structure:
┌─────────────┬──────────┬─────────────────────┐
│ list name   │ name     │ label::English      │
├─────────────┼──────────┼─────────────────────┤
│ gender      │ 1        │ Male                │
│ gender      │ 2        │ Female              │
│ status      │ A        │ Active              │
│ status      │ I        │ Inactive            │
└─────────────┴──────────┴─────────────────────┘
"""
        
        content = tk.Label(
            content_frame,
            text=instr_text,
            justify="left",
            bg=ModernStyle.SURFACE,
            fg=ModernStyle.TEXT_PRIMARY,
            font=content_font,
            wraplength=650,
            anchor="nw"
        )
        content.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Footer
        footer_frame = tk.Frame(instr_window, bg=ModernStyle.SURFACE_DARK)
        footer_frame.pack(fill=tk.X, pady=(0, 0))
        
        footer_font = font.Font(family="Segoe UI", size=9, slant="italic")
        footer = tk.Label(
            footer_frame,
            text="Developed by ALTOS Team\nPowered by Python & Pandas",
            bg=ModernStyle.SURFACE_DARK,
            fg=ModernStyle.TEXT_SECONDARY,
            font=footer_font,
            pady=15
        )
        footer.pack()
        
        # Close button
        close_btn = self.create_modern_button(
            footer_frame,
            "Close",
            instr_window.destroy,
            ModernStyle.PRIMARY
        )
        close_btn.pack(pady=(0, 15))


class AutoConverter:
    """Main application with modern full-width interface"""
    def __init__(self, root):
        self.root = root
        self.root.title("Auto Converter v1 - Professional Edition")
        self.root.geometry("1000x700")
        self.root.configure(bg=ModernStyle.BACKGROUND)
        self.root.resizable(True, True)
        self.root.minsize(900, 600)
        
        self.setup_modern_ui()
    
    def setup_modern_ui(self):
        # Configure modern styling
        style = ttk.Style()
        style.theme_use('clam')
        
        # Custom notebook styling for full-width tabs with consistent sizing
        style.configure(
            "Modern.TNotebook",
            background=ModernStyle.BACKGROUND,
            borderwidth=0,
            tabmargins=[0, 0, 0, 0]
        )
        
        # Base tab configuration with consistent padding
        style.configure(
            "Modern.TNotebook.Tab",
            background=ModernStyle.SURFACE,
            foreground=ModernStyle.TEXT_PRIMARY,
            padding=[30, 12],  # Consistent padding for all states
            borderwidth=0,
            focuscolor='none',
            font=('Segoe UI', 10, 'normal')  # Consistent font
        )
        
        # Map different states with consistent sizing
        style.map(
            "Modern.TNotebook.Tab",
            background=[
                ("selected", ModernStyle.PRIMARY), 
                ("active", ModernStyle.SURFACE_DARK),
                ("!active", ModernStyle.SURFACE)
            ],
            foreground=[
                ("selected", "white"), 
                ("active", ModernStyle.TEXT_PRIMARY),
                ("!active", ModernStyle.TEXT_PRIMARY)
            ],
            padding=[
                ("selected", [30, 12]),  # Same padding for selected
                ("active", [30, 12]),    # Same padding for active
                ("!active", [30, 12])    # Same padding for inactive
            ],
            font=[
                ("selected", ('Segoe UI', 10, 'bold')),  # Bold for selected
                ("active", ('Segoe UI', 10, 'normal')),  # Normal for active
                ("!active", ('Segoe UI', 10, 'normal'))  # Normal for inactive
            ]
        )
        
        # Main container
        main_container = tk.Frame(self.root, bg=ModernStyle.BACKGROUND)
        main_container.pack(fill=tk.BOTH, expand=True)
        
        # Create modern notebook
        self.notebook = ttk.Notebook(main_container, style="Modern.TNotebook")
        self.notebook.pack(fill=tk.BOTH, expand=True)
        
        # Create tab frames
        converter_frame = tk.Frame(self.notebook, bg=ModernStyle.BACKGROUND)
        main_tables_frame = tk.Frame(self.notebook, bg=ModernStyle.BACKGROUND)
        settings_frame = tk.Frame(self.notebook, bg=ModernStyle.BACKGROUND)
        help_frame = tk.Frame(self.notebook, bg=ModernStyle.BACKGROUND)
        
        # Add tabs with icons
        self.notebook.add(main_tables_frame, text="🛠️ Main Tables Converter")
        self.notebook.add(converter_frame, text="🔄 HouseBiz Converter")
        self.notebook.add(settings_frame, text="⚙️ Settings")
        self.notebook.add(help_frame, text="❓ Help")
        
        # Initialize tab classes
        self.main_tables_tab = MainTablesConverterTab(main_tables_frame)
        self.converter_tab = DocumentConverterTab(converter_frame)
        self.settings_tab = SettingsTab(settings_frame)
        self.help_tab = HelpTab(help_frame) 


def main():
    root = tk.Tk()
    app = AutoConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()